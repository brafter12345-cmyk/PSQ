"""document_extract.py — convert an uploaded risk-assessment form to text.

Phase 1 (this module): local, on-prem extraction for the common machine-readable
formats — native (text) PDFs, Word, Excel, plain text. Scanned/image PDFs and images
are detected and flagged for the vision step (Claude PDF/vision), which is wired in Phase 2.

Returns a uniform dict:
  { ok, format, pages, native, chars, needs_ocr, text, blocks, notes }
"""
import re
from pathlib import Path

# Per-PAGE detection so nothing slips through: a genuine text page yields hundreds of characters,
# a scanned/image page yields ~none. Any page under this threshold is routed to OCR.
PAGE_TEXT_MIN = 40

# Scan-seed detection — most of what an external scan needs (domain/website/IP) is already on
# the assessment form. These seeds are extracted LOCALLY and used to trigger the scanner in the
# closed environment; they are still redacted before the mapping LLM call (no conflict).
_EMAIL_DOMAIN = re.compile(r"[\w.+-]+@([a-z0-9.-]+\.[a-z]{2,})", re.I)
_URL = re.compile(r"\b(?:https?://|www\.)([a-z0-9.-]+\.[a-z]{2,})", re.I)
_DOMAIN = re.compile(r"\b((?:[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.)+(?:co\.za|org\.za|net\.za|gov\.za|ac\.za|web\.za|com|net|org|io|biz|info|africa|za))\b", re.I)
_IP = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")
_SKIP_DOMAINS = {"gmail.com", "outlook.com", "hotmail.com", "yahoo.com", "example.com", "google.com"}


def find_scan_seeds(text):
    """Candidate external-scan targets found in the form (broker confirms the primary domain)."""
    text = text or ""
    domains = set()
    for pat in (_EMAIL_DOMAIN, _URL, _DOMAIN):
        for m in pat.findall(text):
            d = m.lower().strip(".")
            if d and d not in _SKIP_DOMAINS and not d.endswith((".png", ".jpg", ".jpeg", ".pdf", ".docx")):
                domains.add(d)
    ips = set(_IP.findall(text))
    ips = {ip for ip in ips if all(0 <= int(o) <= 255 for o in ip.split("."))}
    return {"domains": sorted(domains)[:25], "ips": sorted(ips)[:25]}


def extract(filepath: str, filename: str) -> dict:
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    try:
        if ext == "pdf":
            r = _extract_pdf(filepath)
        elif ext in ("docx",):
            r = _extract_docx(filepath)
        elif ext in ("xlsx", "xlsm"):
            r = _extract_xlsx(filepath)
        elif ext in ("txt", "md", "csv"):
            r = _extract_text(filepath, ext)
        elif ext in ("png", "jpg", "jpeg", "tif", "tiff", "bmp", "webp"):
            r = {"ok": True, "format": "image", "pages": 1, "native": False, "chars": 0,
                 "needs_ocr": True, "text": "", "blocks": [],
                 "notes": ["Image file — text extraction needs the vision/OCR step (Phase 2)."]}
        elif ext == "doc":
            return {"ok": False, "format": "doc", "error": "Legacy .doc not supported — please save as .docx or PDF."}
        else:
            return {"ok": False, "format": ext or "unknown", "error": f"Unsupported file type: .{ext or '?'}"}
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "format": ext, "error": f"Extraction failed: {exc}"}
    # candidate external-scan targets (domains/IPs on the form) — used locally to seed the scanner
    if r.get("ok"):
        r["scan_seeds"] = find_scan_seeds(r.get("text", ""))
    return r


def _extract_pdf(filepath: str) -> dict:
    """Per-page classification: each page is either a text page (extract) or a scanned/image
    page (route to OCR). needs_ocr is True if ANY page needs OCR, so a hybrid document can't be
    silently mapped from its text pages alone."""
    import fitz  # PyMuPDF
    doc = fitz.open(filepath)
    pages = doc.page_count
    blocks, parts, text_pages, ocr_pages = [], [], [], []
    for i, page in enumerate(doc):
        t = (page.get_text("text") or "").strip()
        if len(t) >= PAGE_TEXT_MIN:
            parts.append(t)
            blocks.append({"page": i + 1, "text": t})
            text_pages.append(i + 1)
        else:
            ocr_pages.append(i + 1)  # little/no embedded text -> OCR (err toward OCR so nothing is missed)
    doc.close()
    total = sum(len(p) for p in parts)
    needs_ocr = len(ocr_pages) > 0
    notes = []
    if needs_ocr:
        which = ", ".join(map(str, ocr_pages[:12])) + ("…" if len(ocr_pages) > 12 else "")
        kind = "scanned/image PDF" if not text_pages else "mixed PDF (some text pages, some scanned)"
        notes.append("Detected a %s: %d of %d page(s) have no embedded text and need OCR (page %s). "
                     "Routed to the OCR step so no page's content is missed." % (kind, len(ocr_pages), pages, which))
    return {"ok": True, "format": "pdf", "pages": pages, "native": not needs_ocr,
            "chars": total, "needs_ocr": needs_ocr,
            "text": "\n\n".join(parts), "blocks": blocks,
            "text_pages": text_pages, "ocr_pages": ocr_pages, "notes": notes}


def _extract_docx(filepath: str) -> dict:
    import docx  # python-docx
    d = docx.Document(filepath)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # tables (risk assessments are often tabular Q/A)
    for ti, table in enumerate(d.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("[Table %d]\n%s" % (ti + 1, "\n".join(rows)))
    text = "\n".join(parts)
    return {"ok": True, "format": "docx", "pages": None, "native": True,
            "chars": len(text), "needs_ocr": False, "text": text,
            "blocks": [{"page": None, "text": text}], "notes": []}


def _extract_xlsx(filepath: str) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.append("[Sheet: %s]\n%s" % (ws.title, "\n".join(rows)))
    wb.close()
    text = "\n\n".join(parts)
    return {"ok": True, "format": "xlsx", "pages": len(parts), "native": True,
            "chars": len(text), "needs_ocr": False, "text": text,
            "blocks": [{"page": None, "text": text}], "notes": []}


def _extract_text(filepath: str, ext: str) -> dict:
    text = Path(filepath).read_text(encoding="utf-8", errors="replace")
    return {"ok": True, "format": ext, "pages": None, "native": True,
            "chars": len(text), "needs_ocr": False, "text": text,
            "blocks": [{"page": None, "text": text}], "notes": []}
