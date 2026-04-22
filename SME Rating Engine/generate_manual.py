"""
Generate SME Rating Engine User Manual as a Word document.
Phishield UMA (Pty) Ltd
Enhanced version with flowchart, expanded examples, glossary, market conditions,
common mistakes, version history, and improved screenshot placeholders.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Branding colours
PHISHIELD_DARK = RGBColor(0x0F, 0x17, 0x2A)   # Dark navy
PHISHIELD_ACCENT = RGBColor(0xF9, 0x73, 0x16)  # Orange accent
PHISHIELD_BLUE = RGBColor(0x1E, 0x40, 0xAF)    # Blue
PHISHIELD_GREEN = RGBColor(0x05, 0x96, 0x69)    # Green for tips
PHISHIELD_RED = RGBColor(0xDC, 0x26, 0x26)      # Red for warnings
PHISHIELD_GREY = RGBColor(0x64, 0x74, 0x8B)     # Grey text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = PHISHIELD_DARK
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.3

# ── Heading styles ────────────────────────────────────────────────────────────
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = PHISHIELD_DARK
h1_style.paragraph_format.space_before = Pt(24)
h1_style.paragraph_format.space_after = Pt(8)

h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(13)
h2_style.font.bold = True
h2_style.font.color.rgb = PHISHIELD_DARK
h2_style.paragraph_format.space_before = Pt(16)
h2_style.paragraph_format.space_after = Pt(8)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PHISHIELD_DARK
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
    return h

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def add_para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_tip_box(text, box_type="tip"):
    """Add a styled tip/note/warning paragraph."""
    prefix_map = {
        "tip": ("TIP: ", PHISHIELD_GREEN),
        "note": ("NOTE: ", PHISHIELD_BLUE),
        "warning": ("WARNING: ", PHISHIELD_RED),
        "important": ("IMPORTANT: ", PHISHIELD_ACCENT),
    }
    prefix, color = prefix_map.get(box_type, ("NOTE: ", PHISHIELD_BLUE))
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run_prefix = p.add_run(prefix)
    run_prefix.bold = True
    run_prefix.font.color.rgb = color
    run_prefix.font.size = Pt(10)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = PHISHIELD_GREY
    return p

_figure_counter = [0]  # mutable counter for figure numbering

# Map of old screenshot captions to new figure captions
_figure_captions = {
    "Step 1 — Complete form with all fields filled in, showing turnover info panel and UW outcome badge":
        "Step 1 — Client details form with company name, searchable industry dropdown, turnover inputs, and employee count fields",
    "Step 1 — Underwriting questions with Yes/No toggles and outcome badge":
        "Step 1 — Underwriting Assessment questions with Yes/No toggles and Standard Rates outcome badge",
    "Step 2 — Recommended cover cards with Micro SME badge and FP selector":
        "Step 2 — Recommended cover option cards with Micro SME badge, FP selector, and estimated premium table",
    "Step 3 — Comparison table showing Phishield vs Industry benchmark with delta values":
        "Step 3 — Competitor comparison with quote options auto-matched to selected covers and Industry Benchmark table",
    "Step 4 — Discount inputs with live comparison bars showing green \"competitive\" status":
        "Step 4 — Discount inputs with 10% posture discount applied, comparison bars showing competitive status (green)",
    "Step 5 — Quote summary with audit trail table and export buttons":
        "Step 5 — Quote summary with client details, underwriting outcome, audit trail, and final premium breakdown",
}

def add_screenshot_placeholder(caption):
    _figure_counter[0] += 1
    fig_num = _figure_counter[0]
    # Use mapped caption if available, otherwise use the original
    display_caption = _figure_captions.get(caption, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"Figure {fig_num}: {display_caption}")
    run.italic = True
    run.font.color.rgb = PHISHIELD_GREY
    run.font.size = Pt(10)
    return p

def add_numbered_step(number, text, detail=None):
    p = doc.add_paragraph()
    run_num = p.add_run(f"Step {number}: ")
    run_num.bold = True
    run_num.font.color.rgb = PHISHIELD_ACCENT
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run_detail = p2.add_run(detail)
        run_detail.font.size = Pt(10)
        run_detail.font.color.rgb = PHISHIELD_GREY
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Helper to set cell vertical padding
    def _set_cell_padding(cell, top=60, bottom=60, left=80, right=80):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")} '
            f'w:top="{top}" w:left="{left}" w:bottom="{bottom}" w:right="{right}"/>'
        )
        # Remove existing tcMar if present
        for existing in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(existing)
        tcPr.append(tcMar)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _set_cell_padding(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = WHITE
                run.font.name = 'Calibri'
        set_cell_shading(cell, "0F172A")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            _set_cell_padding(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F1F5F9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def add_flowchart_line(text, indent=0, bold=False, color=None):
    """Add a single line of the text-based flowchart."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = PHISHIELD_DARK
    return p


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Add some blank lines for vertical centring effect
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PHISHIELD")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = PHISHIELD_ACCENT
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("SME Rating Engine")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = PHISHIELD_DARK
run.font.name = 'Calibri'

doc_type = doc.add_paragraph()
doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doc_type.add_run("User Manual")
run.font.size = Pt(20)
run.font.color.rgb = PHISHIELD_GREY
run.font.name = 'Calibri'

doc.add_paragraph()

product = doc.add_paragraph()
product.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = product.add_run("Cyber Protect Business Policy (Risk Rated)")
run.font.size = Pt(12)
run.italic = True
run.font.color.rgb = PHISHIELD_GREY

doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run("Administrator: Phishield UMA (Pty) Ltd")
run.font.size = Pt(11)
run.font.color.rgb = PHISHIELD_DARK

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run("Version 1.2  |  March 2026\nConfidential \u2014 Internal Use Only")
run.font.size = Pt(10)
run.font.color.rgb = PHISHIELD_GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("Table of Contents", level=1)

toc_items = [
    "1. Introduction",
    "   1.4 Decision Flowchart",
    "2. Getting Started",
    "3. Step 1: Client & Industry",
    "   3.7.1 Renewal Details (three required fields)",
    "   3.7.2 Market Conditions Explained (Renewals)",
    "   3.7.3 Premium-drop Protection (Renewals)",
    "4. Step 2: Coverage Recommendations & Selection",
    "5. Step 3: Competitor Quotes & Benchmarking",
    "   5.5 Renewal Behaviour & Insights Banner",
    "6. Step 4: Adjustments & Comparison",
    "7. Step 5: Quote Summary & Export",
    "8. Common Scenarios (Worked Examples)",
    "   8.2 Renewal Quote (normal)",
    "   8.2b Renewal \u2014 Rule I triggered (adjacent target)",
    "   8.2c Renewal \u2014 Rule I triggered (ladder gap)",
    "   8.4 Common Mistakes to Avoid",
    "9. Troubleshooting & FAQ",
    "10. Quick Reference Card",
    "11. Glossary of Key Terms",
    "12. Version History / Changelog",
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_paragraph()
add_tip_box(
    "Screenshots can be found in the manual_screenshots folder alongside this document, "
    "or captured fresh by following the steps in this manual.",
    "note"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1. Introduction", level=1)

doc.add_paragraph(
    "The SME Rating Engine is an internal pricing tool used by Phishield UMA (Pty) Ltd underwriters "
    "to calculate cyber insurance premiums for Small and Medium Enterprise (SME) clients. It produces "
    "quotes for the Cyber Protect Business Policy (Risk Rated), underwritten by Bryte Insurance Company Limited."
)

add_heading("1.1 Purpose", level=2)
doc.add_paragraph(
    "This tool replaces manual spreadsheet-based rating. It automates premium calculations, applies "
    "underwriting rules, and produces professional PDF quote outputs. The engine ensures consistency "
    "across all quotes and provides a full audit trail for compliance."
)

add_heading("1.2 Who Should Use This Manual", level=2)
doc.add_paragraph(
    "This manual is intended for underwriters at Phishield UMA, including junior underwriters who "
    "are learning the rating process. Every field, rule, and calculation is explained in plain language "
    "so that you understand not just what to do, but why each step matters."
)

add_heading("1.3 How the Engine Works (Overview)", level=2)
doc.add_paragraph(
    "The engine follows a 5-step wizard flow. You move through each step sequentially, and the engine "
    "calculates premiums in real time as you enter data. The five steps are:"
)

steps_overview = [
    ("Client & Industry", "Enter company details, turnover, industry, and answer underwriting questions."),
    ("Coverage Selection", "Choose cover limits and Funds Protect (FP) options based on recommendations."),
    ("Benchmarking", "Compare Phishield pricing against industry benchmarks or competitor quotes."),
    ("Adjustments", "Apply posture and discretionary discounts; review live comparison bars."),
    ("Summary & Export", "Review the full audit trail and download PDFs or copy the quote to clipboard."),
]

for i, (title_text, desc) in enumerate(steps_overview, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f"Step {i}: {title_text} \u2014 ")
    run_num.bold = True
    run_desc = p.add_run(desc)

add_tip_box(
    "The engine saves every quote to a backend database automatically. You do not need to manually save your work.",
    "tip"
)

# ══════════════════════════════════════════════════════════════════════════════
# 1.4 DECISION FLOWCHART (NEW)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1.4 Decision Flowchart", level=2)

doc.add_paragraph(
    "The following text-based flowchart summarises the decision path from client data entry through "
    "to the final quote. Use this as a quick reference when assessing a new risk."
)

add_screenshot_placeholder("Decision flowchart \u2014 visual representation of the rating process")

flowchart_lines = [
    ("Client Details \u2192 Industry Check", 0, True, PHISHIELD_DARK),
    ("\u251c\u2500 Healthcare / Public Admin \u2192 REFER FOR UNDERWRITING (blocked)", 1, False, PHISHIELD_RED),
    ("\u251c\u2500 Turnover > R200M \u2192 REFER FOR UNDERWRITING (blocked)", 1, False, PHISHIELD_RED),
    ("\u2514\u2500 All other industries \u2192 Continue", 1, False, PHISHIELD_GREEN),
    ("    \u2193", 1, False, PHISHIELD_GREY),
    ("Underwriting Assessment", 0, True, PHISHIELD_DARK),
    ("\u251c\u2500 Q1 = No \u2192 DECLINE (blocked)", 1, False, PHISHIELD_RED),
    ("\u251c\u2500 Q2\u2013Q6: Count \"No\" answers \u2192 Apply loading (0\u201315%)", 1, False, PHISHIELD_ACCENT),
    ("\u251c\u2500 Q7/Q8 (if FP > R250k): No = Condition of Cover", 1, False, PHISHIELD_ACCENT),
    ("\u2514\u2500 Q9 = No \u2192 REFER (flag, can proceed)", 1, False, PHISHIELD_ACCENT),
    ("    \u2193", 1, False, PHISHIELD_GREY),
    ("Coverage Selection", 0, True, PHISHIELD_DARK),
    ("\u251c\u2500 T/O < R50M + Cover \u2264 R5M \u2192 MICRO SME rates", 1, False, PHISHIELD_GREEN),
    ("\u2514\u2500 Otherwise \u2192 STANDARD formula rates", 1, False, PHISHIELD_BLUE),
    ("    \u2193", 1, False, PHISHIELD_GREY),
    ("Industry Modifier Applied (S&T: 1.35\u20131.63x, Finance: 1.28\u20131.46x)", 0, True, PHISHIELD_DARK),
    ("    \u2193", 1, False, PHISHIELD_GREY),
    ("Benchmarking \u2192 Adjustments \u2192 Summary & Export", 0, True, PHISHIELD_DARK),
]

for text, indent, bold, color in flowchart_lines:
    add_flowchart_line(text, indent, bold, color)

doc.add_paragraph()  # spacing

add_tip_box(
    "Refer this flowchart whenever you are unsure about how the engine will treat a particular client profile. "
    "Each branch is enforced automatically by the engine.",
    "tip"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════

add_heading("2. Getting Started", level=1)

add_heading("2.1 Accessing the Engine", level=2)
doc.add_paragraph(
    "The SME Rating Engine is a web-based application. You access it through your browser."
)

make_table(
    ["Environment", "URL"],
    [
        ["Live (Production)", "https://sme-rating-engine.onrender.com"],
        ["Local Development", "http://localhost:5002"],
    ]
)

add_tip_box(
    "Always use the Live (Production) URL for actual client quotes. The local URL is for testing only.",
    "important"
)

add_heading("2.2 Browser Requirements", level=2)
doc.add_paragraph(
    "The engine works best in modern browsers. Recommended browsers:"
)
for browser in ["Google Chrome (latest version)", "Microsoft Edge (latest version)", "Mozilla Firefox (latest version)"]:
    doc.add_paragraph(browser, style='List Bullet')

add_tip_box(
    "Internet Explorer is not supported. If the interface looks broken, check that you are using a supported browser.",
    "warning"
)

add_heading("2.3 Screen Layout", level=2)
doc.add_paragraph(
    "When you open the engine, you will see:"
)
doc.add_paragraph("Header bar with the Phishield branding", style='List Bullet')
doc.add_paragraph("Progress bar showing the 5 steps (Client, Coverage, Compare, Adjust, Summary)", style='List Bullet')
doc.add_paragraph("The active step panel with form fields", style='List Bullet')
doc.add_paragraph("A floating quote ticker (appears from Step 2 onward) showing the estimated monthly premium", style='List Bullet')

add_screenshot_placeholder("Main screen layout with header and progress bar")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. STEP 1: CLIENT & INDUSTRY
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3. Step 1: Client & Industry", level=1)

doc.add_paragraph(
    "This is the data-capture step. You enter all client information and answer underwriting questions. "
    "The engine uses this data to determine the applicable revenue band, premium rates, and whether the "
    "risk can proceed or must be referred."
)

add_screenshot_placeholder("Step 1 \u2014 Complete form with all fields filled in, showing turnover info panel and UW outcome badge")

# ── 3.1 Company Name ──

add_heading("3.1 Company Name", level=2)
doc.add_paragraph(
    "Enter the full registered name of the client company. This appears on the quote output and PDF."
)
add_tip_box("Use the legal entity name, e.g. 'Acme Trading (Pty) Ltd' rather than just 'Acme'.", "tip")

# ── 3.2 Industry ──

add_heading("3.2 Industry Selection", level=2)
doc.add_paragraph(
    "The industry field is a searchable dropdown. Start typing to filter options, or click the dropdown "
    "arrow to browse. Industries are grouped by main category (e.g., Manufacturing, Services, Retail Trade)."
)
doc.add_paragraph(
    "The industry you select affects the premium in two ways:"
)
doc.add_paragraph("Industry modifier: Software & Technology and Finance industries attract higher premiums (see Section 3.2.1).", style='List Bullet')
doc.add_paragraph("Referral: Healthcare and Public Administration industries are always referred to a senior underwriter and cannot be quoted through the engine.", style='List Bullet')

add_heading("3.2.1 Industry Modifiers", level=3)
doc.add_paragraph(
    "Industry modifiers are multipliers applied to the base premium only (not to the Funds Protect cost). "
    "They reflect the higher cyber risk profile of certain industries."
)

make_table(
    ["Industry", "Band 1-3", "Band 4", "Band 5", "Band 6"],
    [
        ["Software & Technology", "1.35x", "1.41x", "1.51x", "1.63x"],
        ["Finance / Insurance / Real Estate", "1.28x", "1.30x", "1.36x", "1.46x"],
        ["All Other Industries", "1.00x", "1.00x", "1.00x", "1.00x"],
    ]
)

add_tip_box(
    "Finance sub-industries include: Depository Institutions, Non-depository Credit Institutions, "
    "Security & Commodity Brokers, Insurance Carriers, Insurance Agents/Brokers, Real Estate, "
    "and Holding & Other Investment Offices.",
    "note"
)

add_heading("3.2.2 Referred Industries", level=3)
doc.add_paragraph(
    "The following industries always trigger 'Refer for Underwriting' and block the quote:"
)
doc.add_paragraph("Healthcare (all sub-categories: Healthcare Services, Hospitals, Medical Practices, Pharmacies)", style='List Bullet')
doc.add_paragraph("Public Administration (Government Services, Non-classifiable Establishments)", style='List Bullet')

add_tip_box(
    "If a client's industry triggers a referral, contact a senior underwriter. The engine will display "
    "a blocker overlay and the Continue button will be disabled.",
    "warning"
)

# ── 3.3 Turnover ──

add_heading("3.3 Turnover", level=2)
doc.add_paragraph(
    "You must enter two turnover figures:"
)
doc.add_paragraph("Previous Financial Year Turnover \u2014 the completed, audited annual turnover.", style='List Bullet')
doc.add_paragraph("Current Year Estimated Revenue \u2014 the projected revenue for the current financial year.", style='List Bullet')

add_heading("3.3.1 Actual Turnover Calculation", level=3)
doc.add_paragraph(
    "The engine calculates the Actual Turnover as the midpoint (average) of the two figures:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Actual Turnover = (Previous Year + Current Year) / 2")
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph(
    "This midpoint approach is used for conservative underwriting. It smooths out fluctuations and "
    "gives a more stable basis for pricing."
)

add_bold_para("Example:")
doc.add_paragraph(
    "If Previous Year = R12,000,000 and Current Year = R18,000,000, then:\n"
    "Actual Turnover = (R12,000,000 + R18,000,000) / 2 = R15,000,000"
)

add_tip_box(
    "If only one turnover figure is available, the engine uses that single figure as the Actual Turnover.",
    "note"
)

add_heading("3.3.2 Revenue Bands", level=3)
doc.add_paragraph(
    "The Actual Turnover is matched to a Revenue Band. This band determines the base premium rates "
    "and which cover limits are recommended."
)

make_table(
    ["Revenue Band", "Turnover Range"],
    [
        ["Band 1", "R0 \u2013 R10,000,000"],
        ["Band 2", "R10,000,001 \u2013 R25,000,000"],
        ["Band 3", "R25,000,001 \u2013 R50,000,000"],
        ["Band 4", "R50,000,001 \u2013 R75,000,000"],
        ["Band 5", "R75,000,001 \u2013 R100,000,000"],
        ["Band 6", "R100,000,001 \u2013 R150,000,000"],
        ["Band 7", "R150,000,001 \u2013 R200,000,000"],
    ]
)

add_tip_box(
    "Turnover above R200,000,000 cannot be quoted in the engine. The system will display "
    "'Refer for Underwriting' and block further progress. Contact a senior underwriter.",
    "warning"
)

# ── 3.4 Employee Count & Website ──

add_heading("3.4 Employee Count & Website Address", level=2)
doc.add_paragraph(
    "Enter the number of employees and the company's website address. These are recorded for reference "
    "and appear on the quote output. They do not directly affect the premium calculation."
)

# ── 3.5 Underwriting Questions ──

add_heading("3.5 Underwriting Assessment (Q1\u2013Q9)", level=2)
doc.add_paragraph(
    "The underwriting questions assess the client's cyber security posture. Each question has a Yes (Y) "
    "or No (N) toggle. Your answers directly affect whether the quote can proceed, and whether a loading "
    "is applied to the premium."
)

add_heading("3.5.1 Question-by-Question Guide", level=3)

# Q1
add_bold_para("Q1: Active Internet Security Software")
doc.add_paragraph(
    "\"Does your business have an active, comprehensive, paid-for internet security software installed "
    "on all computer systems?\""
)
add_tip_box(
    "Q1 = No is an automatic decline. The client cannot be quoted without basic security software. "
    "This is a hard blocker \u2014 the engine will not allow you to proceed.",
    "warning"
)

# Q2
add_bold_para("Q2: Data Back-Up (Compound Question)")
doc.add_paragraph(
    "This question has two parts that must both be answered:"
)
doc.add_paragraph("Q2.1: \"Do you back up your data on a weekly basis?\"", style='List Bullet')
doc.add_paragraph("Q2.2: \"Do you perform recovery testing at least once per year?\"", style='List Bullet')
doc.add_paragraph(
    "Q2 is considered 'Yes' only when both Q2.1 AND Q2.2 are 'Yes'. If either is 'No', Q2 counts as a "
    "'No' answer for loading purposes."
)

# Q3-Q6
add_bold_para("Q3\u2013Q6: Security Practices")
doc.add_paragraph("Q3: Is data stored separately from the main computer (cloud or offline disk)?")
doc.add_paragraph("Q4: Are computers regularly updated and patched with latest security patches?")
doc.add_paragraph("Q5: Are emails checked for viruses/malware via an email filter?")
doc.add_paragraph("Q6: Are employees regularly advised about secure computer use and internet/email dangers?")

add_heading("3.5.2 Underwriting Loading Rules", level=3)
doc.add_paragraph(
    "The engine counts the number of 'No' answers to Q2\u2013Q6 (treating Q2 as a single question) "
    "and applies loadings as follows:"
)

make_table(
    ["No Answers (Q2\u2013Q6)", "Loading", "Outcome Label"],
    [
        ["0", "0%", "Standard Rates"],
        ["1", "0%", "Proceed with Caution"],
        ["2", "5%", "5% Loading Applied"],
        ["3", "10%", "10% Loading Applied"],
        ["4", "10%", "10% Loading Applied"],
        ["5", "15%", "15% Loading Applied"],
    ]
)

add_tip_box(
    "If all Q2\u2013Q6 are 'No' (5 No answers), the risk is declined outright.",
    "warning"
)

# Q7-Q8
add_bold_para("Q7 & Q8: Funds Protect Questions (Conditional)")
doc.add_paragraph(
    "Q7 and Q8 only appear when the Funds Protect (FP) cover exceeds R250,000. This happens automatically "
    "for cover limits of R7.5M and above (since their base FP starts at R500,000+), or when you manually "
    "select an FP tier above R250,000."
)
doc.add_paragraph(
    "Q7 has three sub-parts asking about documented procedures for:"
)
doc.add_paragraph("Q7.1: Vetting new vendors/customers/payees", style='List Bullet')
doc.add_paragraph("Q7.2: Verifying new beneficiaries on banking profiles", style='List Bullet')
doc.add_paragraph("Q7.3: Verifying requests to amend existing beneficiary payment details", style='List Bullet')

doc.add_paragraph(
    "Q8: \"Do you utilise account verification services offered by your bank or third-party provider?\""
)

add_tip_box(
    "Q7/Q8 = No does NOT trigger a loading or decline. Instead, the relevant item becomes a Condition "
    "of Cover, which is noted on the quote output and PDF. This means the client must implement the "
    "procedure as a condition of their policy.",
    "important"
)

# Q9
add_bold_para("Q9: Prior Cyber Liability Cover")
doc.add_paragraph(
    "\"Have you been covered for cyber liability risks in the last 12 months prior to the inception date "
    "of this policy?\""
)
add_tip_box(
    "When Quote Type = Renewal, Q9 is automatically set to Yes and locked \u2014 an existing Phishield "
    "policy implies prior cover by definition. You cannot change Q9 on a renewal.",
    "important"
)
add_tip_box(
    "On New Business and Competing Quote, Q9 remains editable. Q9 = No on these types is informational "
    "and does not on its own block the quote. However, if Q9 arrives as No on a Renewal via any path "
    "(e.g. data load), the engine treats this as a contradiction and blocks the quote as Refer to "
    "Senior UW.",
    "note"
)

# ── 3.6 Prior Claim ──

add_heading("3.6 Prior Claim", level=2)
doc.add_paragraph(
    "A checkbox at the bottom of Step 1 asks: 'Prior claim in previous term'. Tick this if the client "
    "has had a cyber insurance claim in their previous policy term."
)
add_tip_box(
    "Ticking 'Prior claim' is a hard block. The underwriting outcome is set to Refer to Senior UW, the "
    "'Continue to Coverage' button is disabled, and a blocker overlay is shown. The quote cannot proceed "
    "without senior underwriter review. Untick the box to clear the block if entered in error.",
    "warning"
)

# ── 3.7 Quote Type ──

add_heading("3.7 Quote Type", level=2)
doc.add_paragraph(
    "Select one of three quote types. The type you choose changes the behaviour of later steps."
)

make_table(
    ["Quote Type", "When to Use", "Effect on Engine"],
    [
        ["New Business", "Client has no existing Phishield policy",
         "Standard flow. Industry benchmark (IToo) comparison in Step 3."],
        ["Renewal", "Client has an existing Phishield policy being renewed",
         "Enter current cover limit, current annual premium, and current FP sub-limit. Q9 is auto-set to Yes. Comparison becomes 'Existing Policy' vs new quote. FP equivalent auto-set to Yes. Market condition badge shown. Triggers renewal-specific recommendation logic on Step 2 (Current Cover pin + market-condition options + Premium-drop protection)."],
        ["Competing Quote", "Client has a quote from another insurer",
         "Competitor details entered in Step 3. Comparison against competitor pricing."],
    ]
)

add_heading("3.7.1 Renewal Details", level=3)
doc.add_paragraph(
    "When you select 'Renewal', three additional fields appear, all of which are required before the "
    "'Continue to Coverage' button will enable:"
)
doc.add_paragraph("Current Cover Limit \u2014 select the client's existing cover limit from the dropdown.", style='List Bullet')
doc.add_paragraph("Current Annual Premium \u2014 enter the premium they are currently paying (gross of any discounts already applied).", style='List Bullet')
doc.add_paragraph("Current Funds Protect Sub-limit \u2014 select the FP sub-limit on the existing policy. This is required for apples-to-apples premium comparison by the Premium-drop Protection rule.", style='List Bullet')
add_tip_box(
    "All three fields must be populated. Until they are, 'Continue to Coverage' stays disabled. Switching "
    "away from Renewal (e.g. to New Business) clears all three fields and unlocks Q9.",
    "tip"
)
doc.add_paragraph(
    "A market condition badge is displayed (currently: 'Softening market for 2026'). The market condition "
    "drives which secondary options (upgrades / alternatives / downgrades) appear alongside the existing "
    "cover on Step 2 \u2014 see section 3.7.2."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.7.2 MARKET CONDITIONS EXPLAINED (NEW)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3.7.2 Market Conditions Explained (Renewals)", level=3)

doc.add_paragraph(
    "Understanding the current market condition is critical for renewals. The engine displays a market "
    "condition badge, but it is up to you to apply the right strategy. Below are the three market states "
    "and the recommended approach for each."
)

add_bold_para("Softening Market (Current 2026)")
doc.add_paragraph(
    "Premium rates are decreasing across the industry. Strategy: the engine pins the existing cover as "
    "'Current Cover' and adds up to two higher cover limits as 'Upgrade Option' cards. The client can "
    "retain the same cover for a similar or lower premium, or trade up to higher cover at competitive "
    "rates."
)
add_tip_box(
    "In a softening market, always present at least one upgrade option alongside the like-for-like renewal. "
    "Clients are more receptive to increasing cover when rates are falling.",
    "tip"
)

add_bold_para("Stable Market")
doc.add_paragraph(
    "Rates are steady with minor inflation adjustment. Strategy: the engine pins the existing cover as "
    "'Current Cover' and adds one lower and one higher cover as 'Alternative' cards so the underwriter "
    "can show the price point either side of the current cover. Discretionary discounts should be minimal. "
    "Focus on the value of the Funds Protect benefit as a differentiator."
)

add_bold_para("Hardening Market")
doc.add_paragraph(
    "Rates are increasing due to higher claims frequency or severity. Strategy: the engine pins the "
    "existing cover as 'Current Cover' and adds one lower cover as a 'Downgrade Option' \u2014 giving the "
    "underwriter a conversation-starter if the client is price-sensitive. Be transparent about the "
    "increase, emphasise the value of comprehensive cover, and use the audit trail to justify the premium. "
    "The comparison bars in Step 4 help demonstrate that the increase is market-wide, not Phishield-specific."
)
add_tip_box(
    "In a hardening market, the audit trail is your best tool. Walk the broker through each calculation "
    "step so they can explain the increase to their client.",
    "tip"
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.7.3 PREMIUM-DROP PROTECTION (NEW)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3.7.3 Premium-drop Protection (Renewals)", level=3)

doc.add_paragraph(
    "On renewals, the engine runs an additional check to guard against premium erosion \u2014 particularly "
    "in a softening market where market-rate decreases could silently reduce the client's premium "
    "well below the prior year's. The rule fires regardless of market condition so that any anomalous "
    "drop is flagged."
)

add_bold_para("Trigger")
doc.add_paragraph(
    "When the renewal inputs are all populated, the engine calculates the new premium at the same cover "
    "limit and same FP sub-limit as the existing policy (matching all three inputs the underwriter has "
    "captured). If that new premium is more than 20% lower than the current annual premium, the rule "
    "fires."
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Trigger condition: new premium at same cover/FP < 80% of existing premium")
run.bold = True

add_bold_para("Action when triggered")
doc.add_paragraph(
    "The Step 2 cover recommendations are re-calculated to retain at least 90% of the existing premium. "
    "The card ladder adapts to how far the target sits above the existing cover:"
)
doc.add_paragraph("The existing cover limit is always shown as 'Reference \u2014 Not Recommended' (dimmed, not pre-selected).", style='List Bullet')
doc.add_paragraph("The 'Recommended' target is the lowest cover where the new premium retains at least 90% of the existing.", style='List Bullet')
doc.add_paragraph("When the target is ADJACENT to the existing cover (one step up): 'Reference' + 'Recommended' + one cover above the target as 'Alternative' (natural one-step-further upsell).", style='List Bullet')
doc.add_paragraph("When the target is MORE THAN ONE STEP above the existing cover (i.e. the algorithm skipped covers to reach 90%): every available cover between the existing and target appears as an 'Alternative' card. The card above the target is DROPPED. This prevents ladder gaps that would hide the natural one-step upgrade path.", style='List Bullet')
doc.add_paragraph("Each intermediate 'Alternative' card shows a 'XX% retention' badge \u2014 the new premium expressed as a percentage of the existing premium, at matched FP. The underwriter sees at a glance how close each option sits to the 90% bar, and can decide whether a small discretionary pricing adjustment or an FP tier upgrade is worthwhile to bridge the gap.", style='List Bullet')
doc.add_paragraph("A critical-severity banner appears on Step 2 notifying the underwriter of the potential premium loss and the adjusted recommendation.", style='List Bullet')
doc.add_paragraph("The same information is reflected on the Step 4 UW Conditions Panel and the Step 5 PDF audit trail.", style='List Bullet')

add_tip_box(
    "The retention percentage on intermediate cards is calculated at the MATCHED FP sub-limit \u2014 the "
    "same FP basis used for the Rule I decision. Because higher cover limits have a minimum baseline FP, "
    "an intermediate cover may already use a higher FP than the existing policy (e.g. R7.5M forces a "
    "minimum FP of R500k even if the existing policy was at R250k FP). This is baked into the retention "
    "calculation \u2014 no adjustment required.",
    "note"
)

add_tip_box(
    "Clients are more likely to accept a one-notch upgrade than a two-notch jump. Use the intermediate "
    "Alternative cards to anchor the conversation: 'You're currently at R5M. For similar protection "
    "against premium loss we'd normally recommend R10M, but R7.5M at 83%-retention is a realistic "
    "middle ground if we apply a small discretionary adjustment.'",
    "tip"
)

add_bold_para("Corporate-escalation edge case")
doc.add_paragraph(
    "If the highest available SME cover still produces a premium below 90% retention, the engine "
    "escalates:"
)
doc.add_paragraph("The banner upgrades to 'Premium loss risk \u2014 Corporate referral suggested'.", style='List Bullet')
doc.add_paragraph("The Corporate escalation note is carried onto the Step 4 UW Conditions Panel and the Step 5 PDF audit trail.", style='List Bullet')
doc.add_paragraph("The suggested action is to convert the client to a Corporate policy for higher cover limits \u2014 which requires referral to a senior underwriter.", style='List Bullet')

add_tip_box(
    "The Premium-drop Protection rule does not itself refer the quote to UW \u2014 the underwriter can still "
    "produce the SME quote (at the adjusted target cover, or at the existing cover as reference). The "
    "rule's purpose is to surface the issue so the decision is made consciously. The Corporate-escalation "
    "variant is a stronger hint that the case may have outgrown the SME product.",
    "important"
)

add_bold_para("Revenue band shift")
doc.add_paragraph(
    "A related check: if the existing cover limit is not in the recommended set for the current turnover "
    "band (i.e., revenue has grown or shrunk enough to move the client to a different band), an "
    "informational banner is shown: 'Revenue band shift since last renewal'. This prompts the underwriter "
    "to verify that cover remains adequate. The existing cover is still pinned as 'Current Cover' \u2014 "
    "it is not hidden or de-selected."
)

add_bold_para("UW loading comparison caveat")
doc.add_paragraph(
    "When the current-year underwriting answers produce a loading (Q2\u2013Q6 answers), the new renewal "
    "premium reflects that loading. The prior term's cyber-security posture is not on record in the "
    "rating engine, so a direct year-on-year comparison is not strictly like-for-like. The engine "
    "surfaces this as a note in the Step 2 banner and the Step 4 UW Conditions Panel. Each recommendation "
    "card also shows a small 'UW +X%' badge so the loading is visible at a glance. Industry modifiers "
    "(Software & Technology, Financial) remain baked into the calculation silently \u2014 they are "
    "implied and do not typically change year-on-year."
)

# ── 3.8 Continue ──

add_heading("3.8 Moving to Step 2", level=2)
doc.add_paragraph(
    "The 'Continue to Coverage' button becomes active once all mandatory fields are completed:"
)
doc.add_paragraph("Company Name is entered", style='List Bullet')
doc.add_paragraph("Industry is selected", style='List Bullet')
doc.add_paragraph("At least one turnover figure is entered", style='List Bullet')
doc.add_paragraph("Q1 is answered (Yes or No)", style='List Bullet')
doc.add_paragraph("On Renewal quotes only: Current Cover Limit, Current Annual Premium, AND Current FP Sub-limit are all populated", style='List Bullet')

add_tip_box(
    "If the button remains greyed out, check that all mandatory fields are filled in and that no blockers "
    "are active (Prior claim ticked, Refer for UW, turnover > R200M, Q1 = No, or Renewal-Q9 contradiction).",
    "tip"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. STEP 2: COVERAGE RECOMMENDATIONS & SELECTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading("4. Step 2: Coverage Recommendations & Selection", level=1)

doc.add_paragraph(
    "In this step, the engine recommends appropriate cover limits based on the client's revenue band. "
    "You can accept the recommendations, select custom limits, or build a multi-cover quote with up "
    "to 4 options."
)

add_screenshot_placeholder("Step 2 \u2014 Recommended cover cards with Micro SME badge and FP selector")

add_heading("4.1 Recommended Cover Options", level=2)
doc.add_paragraph(
    "The engine displays cover option cards based on the client's revenue band. Each card shows the "
    "cover limit, annual premium, monthly premium, and included FP cost. Cards are tagged as:"
)

make_table(
    ["Tag", "Meaning"],
    [
        ["Recommended", "Best fit for the client's revenue band (New Business / Competing Quote)"],
        ["Optional", "Available but not the primary recommendation"],
        ["Request Only", "Available on request; needs justification"],
        ["N/A", "Not available for this revenue band"],
        ["Current Cover", "Renewal only \u2014 the client's existing cover limit, pinned on Step 2 even if outside the recommended set for the current band. Auto-selected unless Premium-drop Protection fires."],
        ["Current Cover \u2022 Recommended", "Renewal only \u2014 the existing cover is also in the recommended set for the current band (ideal case)."],
        ["Upgrade Option", "Renewal, softening market \u2014 higher cover shown alongside the current cover."],
        ["Alternative", "Renewal, stable market \u2014 one cover above and one below the current cover, shown for comparison."],
        ["Downgrade Option", "Renewal, hardening market \u2014 one cover below the current cover, offered as a price-sensitive alternative."],
        ["Reference \u2014 Not Recommended", "Renewal, Premium-drop Protection triggered \u2014 the existing cover shown for reference only (dimmed, not auto-selected). The Recommended target appears alongside as a higher cover."],
        ["Alternative + XX% retention", "Renewal, Premium-drop Protection triggered with a ladder gap (target > existing+1). Each skipped cover between the existing and the Recommended target is surfaced as an Alternative, tagged with a 'XX% retention' badge showing how close that cover's premium sits to the 90% retention target at matched FP."],
    ]
)

add_heading("4.1.1 Cover Availability by Revenue Band", level=3)
doc.add_paragraph(
    "The following table shows which cover limits are recommended (R), optional (O), request-only (Q), "
    "or not available (\u2014) for each revenue band:"
)

make_table(
    ["Revenue Band", "R1M", "R2.5M", "R5M", "R7.5M", "R10M", "R15M"],
    [
        ["R0\u2013R10M",     "R", "R", "O", "O", "O", "O"],
        ["R10M\u2013R25M",   "O", "R", "R", "O", "O", "O"],
        ["R25M\u2013R50M",   "O", "R", "R", "O", "O", "O"],
        ["R50M\u2013R75M",   "Q", "O", "R", "R", "O", "O"],
        ["R75M\u2013R100M",  "\u2014", "O", "O", "R", "R", "O"],
        ["R100M\u2013R150M", "\u2014", "\u2014", "O", "R", "R", "O"],
        ["R150M\u2013R200M", "\u2014", "\u2014", "O", "O", "R", "R"],
    ]
)

add_heading("4.2 Selecting Covers", level=2)
doc.add_paragraph(
    "Click on a recommended card to select it. The card will highlight and the premium display updates. "
    "Click again to deselect."
)
doc.add_paragraph(
    "To select a non-recommended limit, click 'Select Custom Cover Limit' to reveal all available options."
)

add_heading("4.3 Multi-Cover Quoting", level=2)
doc.add_paragraph(
    "You can select multiple cover limits to present the client with options. Select 2\u20134 covers by "
    "clicking multiple cards. Each selected cover becomes a separate quote option."
)
doc.add_paragraph("Each option gets its own Funds Protect selection, discounts, and comparison.", style='List Bullet')
doc.add_paragraph("You can add the same cover limit twice with different FP tiers using the '+' button on a card.", style='List Bullet')
doc.add_paragraph("Maximum of 4 options per quote.", style='List Bullet')
doc.add_paragraph("Each option receives a unique quote reference number.", style='List Bullet')

add_heading("4.4 Micro SME Pricing", level=2)
doc.add_paragraph(
    "Micro SME rates apply when both conditions are met:"
)
doc.add_paragraph("Actual Turnover is below R50,000,000 (revenue bands 1\u20133)", style='List Bullet')
doc.add_paragraph("Cover limit is R5M or below (R1M, R2.5M, or R5M)", style='List Bullet')

doc.add_paragraph(
    "When Micro SME applies, a badge is shown: 'Micro SME Rates Applied'. These are flat-rate premiums "
    "that are lower than the standard formula-based rates. Industry modifiers for Software & Technology "
    "and Finance still apply to Micro SME base premiums."
)

make_table(
    ["Cover Limit", "Base Premium", "Base FP Cost", "Total (Base + FP)"],
    [
        ["R1M", "R4,164", "R2,100", "R6,264"],
        ["R2.5M", "R5,904", "R2,616", "R8,520"],
        ["R5M", "R7,908", "R3,144", "R11,052"],
    ]
)

add_heading("4.5 Funds Protect (FP) Cover", level=2)
doc.add_paragraph(
    "Funds Protect is an add-on that covers the client against financial losses from cyber-enabled fraud "
    "(e.g., business email compromise, social engineering). Every quote includes a minimum FP level "
    "based on the cover limit."
)

make_table(
    ["Cover Limit", "Minimum FP", "Minimum FP Cost/yr"],
    [
        ["R1M", "R150,000", "R2,100"],
        ["R2.5M", "R200,000", "R2,616"],
        ["R5M", "R250,000", "R3,144"],
        ["R7.5M", "R500,000", "R4,800"],
        ["R10M", "R1,000,000", "R9,108"],
        ["R15M", "R1,500,000", "R15,756"],
    ]
)

doc.add_paragraph(
    "You can upgrade the FP level by selecting a higher tier from the FP selector cards. Available "
    "upgrade tiers range up to R5M FP (R71,160/yr)."
)

add_tip_box(
    "When FP exceeds R250,000, questions Q7 and Q8 become active in Step 1. If you go back and change "
    "the FP level, the underwriting questions adjust accordingly.",
    "important"
)

# ── 4.6 Excess ──
add_heading("4.6 Excess (Deductible)", level=2)
doc.add_paragraph("Each cover limit has a standard excess amount:")

make_table(
    ["Cover Limit", "Excess"],
    [
        ["R1M", "R10,000"],
        ["R2.5M", "R10,000"],
        ["R5M", "R10,000"],
        ["R7.5M", "R15,000"],
        ["R10M", "R20,000"],
        ["R15M", "R25,000"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. STEP 3: COMPETITOR QUOTES & BENCHMARKING
# ══════════════════════════════════════════════════════════════════════════════

add_heading("5. Step 3: Competitor Quotes & Benchmarking", level=1)

doc.add_paragraph(
    "This step lets you compare Phishield pricing against industry benchmarks (IToo pricing) or "
    "against an actual competitor quote the client has received."
)

add_screenshot_placeholder("Step 3 \u2014 Comparison table showing Phishield vs Industry benchmark with delta values")

add_heading("5.1 Industry Benchmark (Default)", level=2)
doc.add_paragraph(
    "By default, the engine compares your quote against IToo pricing, which serves as the industry "
    "benchmark. IToo does not include a Funds Protect equivalent, so the comparison is structured "
    "to be fair:"
)
doc.add_paragraph("Phishield (with FP): The full premium including Funds Protect.", style='List Bullet')
doc.add_paragraph("Phishield (ex-FP): The premium excluding Funds Protect.", style='List Bullet')
doc.add_paragraph("Industry Benchmark: The IToo premium for the same cover limit and turnover band.", style='List Bullet')
doc.add_paragraph("Delta: The difference between Phishield (ex-FP) and the benchmark.", style='List Bullet')

add_heading("5.2 Competitor Comparison", level=2)
doc.add_paragraph(
    "If the client has an existing quote from another insurer, toggle 'Does the client have existing "
    "quote(s) for comparison?' to 'Yes'. Then enter:"
)
doc.add_paragraph("Competitor / Provider Name (e.g., Guardrisk, Chubb)", style='List Bullet')
doc.add_paragraph("Number of cover limits to compare (1\u20134)", style='List Bullet')
doc.add_paragraph("For each option: Requested Cover Limit, Competitor Overall Limit, and Competitor Premium", style='List Bullet')

add_heading("5.3 FP Equivalent Toggle", level=2)
doc.add_paragraph(
    "An important toggle asks: 'Does the competitor quote include a FP equivalent?'"
)
doc.add_paragraph("If No (default): The comparison uses Phishield ex-FP vs the competitor/benchmark, because comparing a premium that includes FP against one that does not would be unfair.", style='List Bullet')
doc.add_paragraph("If Yes: The comparison uses Phishield full premium (with FP) vs the competitor/benchmark.", style='List Bullet')

add_tip_box(
    "For renewal quotes, the FP equivalent is automatically set to 'Yes' because the existing policy "
    "premium already includes FP-equivalent cover.",
    "note"
)

add_heading("5.4 Understanding the Comparison Table", level=2)
doc.add_paragraph(
    "The comparison table at the bottom of Step 3 shows a row for each selected cover option. "
    "The Delta column is colour-coded:"
)
doc.add_paragraph("Green: Phishield is cheaper than the benchmark (competitive).", style='List Bullet')
doc.add_paragraph("Amber: Phishield is within a small margin of the benchmark.", style='List Bullet')
doc.add_paragraph("Red: Phishield is more expensive than the benchmark.", style='List Bullet')

add_tip_box(
    "A red delta does not mean you cannot quote. It means you may need to apply discounts in Step 4 "
    "to become competitive, or explain the value of the FP cover to the client.",
    "tip"
)

add_heading("5.5 Renewal Behaviour", level=2)
doc.add_paragraph(
    "When the quote type is 'Renewal', the benchmark column changes from 'Industry Benchmark' to "
    "'Existing Policy'. The comparison is then between the new Phishield quote and what the client "
    "is currently paying."
)
doc.add_paragraph(
    "Renewals also show a renewal-insights banner above the comparison table summarising any of the "
    "following that apply:"
)
doc.add_paragraph("Premium-drop Protection triggered (critical severity) \u2014 see section 3.7.3.", style='List Bullet')
doc.add_paragraph("Corporate escalation (critical severity) \u2014 max SME cover still below 90% retention.", style='List Bullet')
doc.add_paragraph("Revenue band shift (informational) \u2014 existing cover outside current recommended set.", style='List Bullet')
doc.add_paragraph("UW loading comparison caveat (informational) \u2014 current-year Q2\u2013Q6 answers produce a loading; prior posture not on record.", style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. STEP 4: ADJUSTMENTS & COMPARISON
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. Step 4: Adjustments & Comparison", level=1)

doc.add_paragraph(
    "In this step you can apply discounts, set a manual premium override, add endorsements, and see "
    "a live visual comparison of your quote against the benchmark."
)

add_screenshot_placeholder("Step 4 \u2014 Discount inputs with live comparison bars showing green \"competitive\" status")

add_heading("6.1 Posture Discount", level=2)
doc.add_paragraph(
    "The Posture Discount reflects the quality of the client's cyber security posture based on your "
    "assessment. Enter a percentage from 0% to 35%."
)
doc.add_paragraph(
    "Use this when the client has demonstrably strong security practices beyond the minimum requirements "
    "(e.g., dedicated CISO, SOC 2 certification, penetration testing programme, security awareness training)."
)

add_heading("6.2 Discretionary Discount", level=2)
doc.add_paragraph(
    "The Discretionary Discount is used for competitive positioning. Enter a percentage from 0% to 35%."
)
doc.add_paragraph(
    "Use this when you need to match or beat a competitor's premium to win the business. This should "
    "be justified and documented."
)

add_heading("6.3 Discount Limits and Warnings", level=2)
doc.add_paragraph(
    "Discounts are applied to the full premium (base + FP). They compound multiplicatively:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Final Premium = Premium x (1 - Posture%) x (1 - Discretionary%)")
run.bold = True

doc.add_paragraph(
    "For example, a 20% posture discount and 10% discretionary discount results in an effective "
    "discount of 28% (not 30%)."
)

add_tip_box(
    "If the combined effective discount exceeds 35%, a warning is displayed: 'Combined discount exceeds "
    "35%. Senior underwriter approval required.' You can still proceed, but the quote will require "
    "sign-off from a senior underwriter.",
    "warning"
)

add_heading("6.4 Manual Premium Override", level=2)
doc.add_paragraph(
    "An optional field allows you to enter a specific premium amount, bypassing the calculated premium. "
    "Use this only in exceptional circumstances where the standard calculation does not apply "
    "(e.g., previously agreed renewal rates)."
)

add_heading("6.5 Underwriting Conditions Summary", level=2)
doc.add_paragraph(
    "If any conditions of cover were triggered by Q7/Q8 answers, they are summarised here. These "
    "conditions will appear on the quote output and PDF."
)

add_heading("6.6 Endorsements / Underwriter Notes", level=2)
doc.add_paragraph(
    "A free-text area for recording endorsements, special conditions, or any notes to be included "
    "on the quote output. Common uses:"
)
doc.add_paragraph("Special terms agreed with the client or broker", style='List Bullet')
doc.add_paragraph("Exclusions specific to this risk", style='List Bullet')
doc.add_paragraph("Notes about the client's risk profile", style='List Bullet')

add_heading("6.7 Live Comparison Bars", level=2)
doc.add_paragraph(
    "The bottom of Step 4 shows a visual bar chart comparing your final premium against the selected "
    "benchmark (Industry or Competitor). You can toggle between 'Industry' and 'Competitor Quote(s)'."
)
doc.add_paragraph(
    "Each bar shows the Phishield premium as a filled bar, with a line marker indicating the benchmark. "
    "The delta (difference) is displayed as both a Rand amount and a percentage."
)

add_heading("6.8 Multi-Cover: Apply to All", level=2)
doc.add_paragraph(
    "When you have multiple quote options, a checkbox appears: 'Apply discounts to all quote options'. "
    "When ticked, the posture and discretionary discounts you enter are applied uniformly to all options. "
    "When unticked, you can set different discounts per option using the option tabs."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. STEP 5: QUOTE SUMMARY & EXPORT
# ══════════════════════════════════════════════════════════════════════════════

add_heading("7. Step 5: Quote Summary & Export", level=1)

doc.add_paragraph(
    "The final step displays a comprehensive summary of the entire quote, including an audit trail "
    "of how the premium was calculated. This is also where you export the quote."
)

add_screenshot_placeholder("Step 5 \u2014 Quote summary with audit trail table and export buttons")

add_heading("7.1 Quote Reference Number", level=2)
doc.add_paragraph(
    "Each quote receives a unique reference number in the format:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CPB-YYYYMMDD-NNNN-[CoverLimit]-FP[FPLimit]")
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph(
    "Where CPB = Cyber Protect Business, YYYYMMDD = date, NNNN = sequential number, "
    "CoverLimit = e.g. R10M, and FPLimit = e.g. R1M."
)
doc.add_paragraph(
    "For multi-cover quotes, all options share the same base reference (CPB-YYYYMMDD-NNNN) "
    "but each has a unique suffix for its cover limit and FP selection."
)

add_heading("7.2 Client Details Summary", level=2)
doc.add_paragraph(
    "A summary card displays: Company Name, Industry, Actual Turnover, Revenue Bracket, Website, "
    "and Quote Type. Review these for accuracy before exporting."
)

add_heading("7.3 Underwriting Summary", level=2)
doc.add_paragraph(
    "Shows the underwriting outcome (Standard Rates / Proceed with Caution / Loading / Refer / Decline), "
    "any loading percentage applied, and conditions of cover."
)

add_heading("7.4 Premium Breakdown (Audit Trail)", level=2)
doc.add_paragraph(
    "For each cover option quoted, a detailed breakdown table shows every calculation step:"
)

make_table(
    ["Step", "Description", "Example"],
    [
        ["1", "Base premium (from table or formula)", "R11,160"],
        ["2", "Industry modifier applied", "R11,160 x 1.00 = R11,160"],
        ["3", "UW loading (if applicable)", "R11,160 x 1.05 = R11,718"],
        ["4", "Funds Protect cost added", "+ R3,144"],
        ["5", "Total before discounts", "R14,862"],
        ["6", "Discounts applied", "x 0.80 = R11,890"],
    ]
)

doc.add_paragraph(
    "Below the breakdown, final figures are displayed: Annual (with FP), Annual (excl FP), and Monthly."
)

add_heading("7.5 Export Options", level=2)
doc.add_paragraph("Three export options are available:")

add_bold_para("Download PDF")
doc.add_paragraph(
    "Generates a professional PDF document for each cover option. The PDF includes all client details, "
    "underwriting summary, premium breakdown, conditions of cover, and Phishield branding. PDFs are "
    "also saved to the backend database for record-keeping."
)

add_bold_para("Copy to Clipboard")
doc.add_paragraph(
    "Copies a plain-text summary of the quote to your clipboard. Useful for pasting into emails or "
    "other communication tools."
)

add_bold_para("Print Quote")
doc.add_paragraph(
    "Opens the browser's print dialog, allowing you to print the quote summary directly or save it "
    "as a PDF using the browser's built-in PDF printer."
)

add_tip_box(
    "All quotes are automatically saved to the backend database when you reach Step 5. You do not need "
    "to manually save. The quote reference number is your key for retrieving the quote later.",
    "tip"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. COMMON SCENARIOS (WORKED EXAMPLES)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("8. Common Scenarios (Worked Examples)", level=1)

# ── Scenario 1: New Business ──

add_heading("8.1 Scenario 1: New Business Quote", level=2)

add_bold_para("Client: TechStart Solutions (Pty) Ltd")
doc.add_paragraph(
    "Industry: Software and Technology\n"
    "Previous Turnover: R8,000,000\n"
    "Current Estimated: R12,000,000\n"
    "Employees: 35\n"
    "All underwriting questions: Yes"
)

add_heading("Walkthrough:", level=3)
add_numbered_step(1, "Enter client details in Step 1.",
    "Actual Turnover = (R8M + R12M) / 2 = R10M. Revenue Band: R0\u2013R10M (Band 1). "
    "Quote Type: New Business.")
add_numbered_step(2, "In Step 2, the engine recommends R1M and R2.5M covers.",
    "Both qualify for Micro SME rates (turnover < R50M, cover <= R5M). "
    "Select both recommended options. "
    "R1M premium: R4,164 base x 1.35 (S&T modifier) = R5,621 + R2,100 FP = R7,721/yr. "
    "R2.5M premium: R5,904 x 1.35 = R7,970 + R2,616 FP = R10,586/yr.")

add_heading("Step 3 Detail: Comparison Table", level=4)
doc.add_paragraph(
    "The comparison table in Step 3 shows the following values for each selected cover option:"
)
make_table(
    ["Cover", "Phishield (with FP)", "Phishield (ex-FP)", "Industry Benchmark (IToo)", "Delta"],
    [
        ["R1M", "R7,721", "R5,621", "R5,375", "+R246 (4.6%)"],
        ["R2.5M", "R10,586", "R7,970", "R7,850", "+R120 (1.5%)"],
    ]
)
doc.add_paragraph(
    "Both options show amber/slight red deltas \u2014 Phishield ex-FP is marginally above the IToo benchmark. "
    "However, the FP cover provides significant additional value that the benchmark does not include."
)

add_heading("Step 4 Detail: Discount Calculation", level=4)
doc.add_paragraph(
    "Client has SOC 2 certification, so apply a 10% posture discount. No discretionary discount needed."
)
make_table(
    ["Cover", "Pre-Discount", "Posture (10%)", "Final Annual", "Final Monthly"],
    [
        ["R1M", "R7,721", "R7,721 x (1 - 0.10) = R6,949", "R6,949", "R579"],
        ["R2.5M", "R10,586", "R10,586 x (1 - 0.10) = R9,527", "R9,527", "R794"],
    ]
)
doc.add_paragraph(
    "The comparison bars now show green for R1M (R6,949 vs R5,375 benchmark \u2014 Phishield ex-FP after "
    "discount = R5,059, now below benchmark) and green for R2.5M (ex-FP after discount = R7,173, "
    "below the R7,850 benchmark)."
)

add_numbered_step(5, "Download PDFs in Step 5 for both options.",
    "Quote refs: CPB-20260324-0001-R1M-FPR150k and CPB-20260324-0001-R2.5M-FPR200k.")

# ── Scenario 2: Renewal ──

add_heading("8.2 Scenario 2: Renewal Quote (normal)", level=2)

add_bold_para("Client: SecureFinance Holdings (Pty) Ltd")
doc.add_paragraph(
    "Industry: Insurance Agents, Brokers, and Service\n"
    "Previous Turnover: R60,000,000\n"
    "Current Estimated: R65,000,000\n"
    "Current Cover: R7.5M at R32,000/yr, FP sub-limit R500,000\n"
    "All underwriting questions: Yes"
)

add_heading("Walkthrough:", level=3)
add_numbered_step(1, "Enter details and select 'Renewal' as Quote Type.",
    "Actual Turnover = (R60M + R65M) / 2 = R62.5M. Revenue Band: R50M\u2013R75M (Band 4). "
    "Enter Current Cover Limit: R7.5M, Current Premium: R32,000, Current FP Sub-limit: R500,000. "
    "Q9 is auto-set to Yes and locked. Market condition: Softening 2026.")
add_numbered_step(2, "Engine pins R7.5M as 'Current Cover \u2022 Recommended' and adds R10M + R15M as Upgrade Options.",
    "The new premium at R7.5M/R500k FP is around R28,450 \u2014 that is only an 11% drop from R32,000, "
    "so Premium-drop Protection does not trigger. R7.5M auto-selects. The underwriter can also add "
    "R10M from the Upgrade Option cards to present an upsell path.")

add_heading("Step 3 Detail: Comparison Table (Renewal)", level=4)
doc.add_paragraph(
    "For renewals, the comparison shows 'Existing Policy' instead of 'Industry Benchmark'. "
    "FP equivalent is Yes, so full premium (with FP) is compared:"
)
make_table(
    ["Cover", "Phishield (with FP)", "Existing Policy", "Delta"],
    [
        ["R7.5M (like-for-like)", "R28,450", "R32,000", "-R3,550 (-11.1%)"],
        ["R10M (upgrade)", "R35,200", "R32,000", "+R3,200 (+10.0%)"],
    ]
)
doc.add_paragraph(
    "The R7.5M renewal shows a green delta \u2014 the client saves R3,550 on a like-for-like renewal "
    "in the softening market. The R10M upgrade is only R3,200 more than their current premium, "
    "giving them 33% more cover for a small increase."
)

add_heading("Step 4 Detail: Discount Calculation (Renewal)", level=4)
doc.add_paragraph(
    "Apply a 5% discretionary discount to sweeten the renewal offer:"
)
make_table(
    ["Cover", "Pre-Discount", "Discretionary (5%)", "Final Annual", "Final Monthly"],
    [
        ["R7.5M", "R28,450", "R28,450 x (1 - 0.05) = R27,028", "R27,028", "R2,252"],
        ["R10M", "R35,200", "R35,200 x (1 - 0.05) = R33,440", "R33,440", "R2,787"],
    ]
)
doc.add_paragraph(
    "After discount, R7.5M is now R4,972 below the existing premium (-15.5%). "
    "R10M at R33,440 is now only R1,440 above the existing premium (+4.5%) \u2014 a compelling "
    "upgrade pitch. The comparison bars show deep green for both."
)

add_numbered_step(5, "Export both options for the broker.",
    "The renewal badge, market condition, and existing FP sub-limit are shown on the PDF.")

# ── Scenario 2b: Renewal with Premium-drop Protection trigger ──

add_heading("8.2b Scenario 2b: Renewal \u2014 Premium-drop Protection triggered (adjacent target)", level=2)

add_bold_para("Client: Acme Transport (Pty) Ltd")
doc.add_paragraph(
    "Industry: Railroad Transportation\n"
    "Previous Turnover: R15,000,000\n"
    "Current Estimated: R18,000,000\n"
    "Current Cover: R5M at R18,000/yr, FP sub-limit R250,000\n"
    "All underwriting questions: Yes"
)

add_heading("Walkthrough:", level=3)
add_numbered_step(1, "Enter details and select 'Renewal'. Fill all three renewal fields.",
    "Current Cover Limit: R5M, Current Premium: R18,000, Current FP Sub-limit: R250,000.")
add_numbered_step(2, "Click Continue \u2014 Step 2 fires the Premium-drop Protection rule.",
    "New premium at R5M/R250k FP calculates to approximately R11,052. That is ~39% below "
    "existing (R18,000) \u2014 past the 20% threshold \u2014 so the rule fires. The algorithm lands on "
    "R7.5M as the lowest cover reaching 90% retention. R7.5M is ADJACENT to R5M so the standard "
    "Reference + Recommended + one-above-target Alternative card set is used:")
doc.add_paragraph("R5M as 'Reference \u2014 Not Recommended' (dimmed, not pre-selected)", style='List Bullet')
doc.add_paragraph("R7.5M as 'Recommended' (lowest cover where new premium \u2265 90% of R18,000) \u2014 auto-selected", style='List Bullet')
doc.add_paragraph("R10M as 'Alternative' (one step further upsell)", style='List Bullet')
doc.add_paragraph("Critical-severity banner summarising the premium drop and the adjusted recommendation", style='List Bullet')

add_tip_box(
    "The auto-selection moves from R5M to R7.5M so the underwriter starts from the protective default. "
    "If the underwriter disagrees they can click R5M (Reference) to add it to the quote anyway \u2014 the "
    "rule surfaces the decision, it does not enforce it.",
    "note"
)

add_numbered_step(3, "The Step 4 UW Conditions Panel and Step 5 PDF audit trail both include the 'Premium Loss Risk on Renewal' note.",
    "This creates a written record that the adjustment was considered \u2014 useful for compliance and "
    "broker conversations.")

# ── Scenario 2c: Rule I with a ladder gap ──

add_heading("8.2c Scenario 2c: Renewal \u2014 Premium-drop Protection triggered (ladder gap)", level=2)

add_bold_para("Client: CapitalStream Real Estate Finance (Pty) Ltd")
doc.add_paragraph(
    "Industry: Finance, Insurance and Real Estate \u2014 Real Estate\n"
    "Previous Turnover: R65,000,000\n"
    "Current Estimated: R75,000,000\n"
    "Current Cover: R5M at R45,000/yr, FP sub-limit R250,000\n"
    "Q3 = No, Q4 = No (5% UW loading applied)"
)

add_heading("Walkthrough:", level=3)
add_numbered_step(1, "Enter details and select 'Renewal'. Fill all three renewal fields.",
    "Current Cover: R5M, Current Premium: R45,000, Current FP Sub-limit: R250,000. "
    "Turnover R70M lands in Band 4 (R50M\u2013R75M). Finance industry modifier applies.")
add_numbered_step(2, "Click Continue \u2014 the rule fires, and the target is more than one step above existing.",
    "New premium at R5M/R250k FP is approximately R29,800 (about 34% below R45,000). The "
    "algorithm searches upward: R7.5M at matched FP (auto-bumped to R500k minimum) produces "
    "approximately R37,500 \u2014 that's 83% retention, still below the 90% bar. R10M at matched FP "
    "(R500k) produces approximately R43,000 \u2014 \u226590% retention, so R10M becomes the Recommended "
    "target. Because the target (R10M) is TWO steps above the existing (R5M), the ladder-gap "
    "logic kicks in and the card set becomes:")
doc.add_paragraph("R5M as 'Reference \u2014 Not Recommended'", style='List Bullet')
doc.add_paragraph("R7.5M as 'Alternative' with a '83% retention' badge (the intermediate rung)", style='List Bullet')
doc.add_paragraph("R10M as 'Recommended' \u2014 auto-selected", style='List Bullet')
doc.add_paragraph("R15M is NOT shown \u2014 the above-target Alternative is dropped when there are intermediate Alternatives", style='List Bullet')

add_tip_box(
    "R7.5M is the natural conversation rung: the client will often prefer a one-notch upgrade over "
    "a two-notch jump. The 83% retention badge tells the underwriter exactly how close R7.5M sits "
    "to the 90% bar \u2014 a small discretionary adjustment on Step 4, or an FP tier bump (e.g. R500k "
    "\u2192 R1M), can often close that 7-percentage-point gap. The Recommended R10M remains the "
    "default so the protective option is always pre-selected.",
    "tip"
)

add_bold_para("Variant: Corporate escalation with full ladder visible")
doc.add_paragraph(
    "If the existing premium was so high that even the highest available SME cover (R15M) still fell "
    "below 90% retention, the banner upgrades to 'Premium loss risk \u2014 Corporate referral "
    "suggested'. The ladder-gap logic still applies \u2014 every intermediate cover between the existing "
    "and R15M (the max-available target) appears as an Alternative with its retention badge. This "
    "gives the underwriter a complete picture of how SME cover falls short and strengthens the case "
    "for conversion to a Corporate product (requires senior-underwriter referral)."
)

# ── Scenario 3: Competing Quote ──

add_heading("8.3 Scenario 3: Competing Quote", level=2)

add_bold_para("Client: GreenRetail (Pty) Ltd")
doc.add_paragraph(
    "Industry: General Merchandise Stores\n"
    "Previous Turnover: R30,000,000\n"
    "Current Estimated: R35,000,000\n"
    "Competitor: Guardrisk at R18,500 for R5M cover (no FP equivalent)\n"
    "Q3 = No, Q6 = No (2 'No' answers)"
)

add_heading("Walkthrough:", level=3)
add_numbered_step(1, "Enter details. Select 'Competing Quote' as Quote Type.",
    "Actual Turnover = (R30M + R35M) / 2 = R32.5M. Revenue Band: R25M\u2013R50M (Band 3). "
    "Q3=No and Q6=No: 2 'No' answers = 5% loading applied. "
    "Industry modifier: 1.00 (Retail Trade).")
add_numbered_step(2, "Select R5M cover. Micro SME applies (turnover < R50M, cover <= R5M).",
    "Base premium: R7,908 x 1.00 (modifier) x 1.05 (loading) = R8,303 + R3,144 FP = R11,447/yr.")

add_heading("Step 3 Detail: Comparison Table (Competing Quote)", level=4)
doc.add_paragraph(
    "Competitor details entered. FP equivalent = No, so Phishield ex-FP is used for comparison:"
)
make_table(
    ["Cover", "Phishield (with FP)", "Phishield (ex-FP)", "Competitor (Guardrisk)", "Delta"],
    [
        ["R5M", "R11,447", "R8,303", "R18,500", "-R10,197 (-55.1%)"],
    ]
)
doc.add_paragraph(
    "Phishield is dramatically cheaper. Even with FP included (R11,447), Phishield is R7,053 below "
    "Guardrisk. The comparison bar will show deep green."
)

add_heading("Step 4 Detail: Discount Calculation (Competing Quote)", level=4)
doc.add_paragraph(
    "No discounts needed \u2014 Phishield is already significantly cheaper. The comparison bar shows:"
)
make_table(
    ["Metric", "Value"],
    [
        ["Phishield Final (with FP)", "R11,447"],
        ["Phishield Final (ex-FP)", "R8,303"],
        ["Competitor", "R18,500"],
        ["Delta (ex-FP vs Competitor)", "-R10,197 (-55.1%)"],
        ["Comparison Bar", "Deep green \u2014 very competitive"],
    ]
)
doc.add_paragraph(
    "Add an endorsement noting the 5% UW loading due to missing data separation (Q3) and lack of "
    "employee security training (Q6). Recommend the client addresses these for future renewal."
)

add_numbered_step(5, "Export the quote. Highlight the competitive advantage in client communication.")

# ══════════════════════════════════════════════════════════════════════════════
# 8.4 COMMON MISTAKES TO AVOID (NEW)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("8.4 Common Mistakes to Avoid", level=2)

doc.add_paragraph(
    "The following are common errors made by underwriters when using the SME Rating Engine. "
    "Review this checklist regularly to avoid costly mistakes."
)

mistakes = [
    (
        "Forgetting to check FP-dependent questions (Q7/Q8) when upgrading FP above R250k",
        "When you upgrade a client's FP above R250,000 in Step 2, Q7 and Q8 become active in Step 1. "
        "If you do not go back to answer them, the engine may not correctly apply conditions of cover. "
        "Always check the underwriting panel after changing FP."
    ),
    (
        "Applying combined discounts above 35% without senior approval",
        "A combined effective discount above 35% triggers a warning. Do not proceed without getting "
        "senior underwriter sign-off first. Document the approval before exporting the quote."
    ),
    (
        "Not entering both turnover figures (using only one gives an inaccurate midpoint)",
        "The engine uses the midpoint of Previous Year and Current Year turnover. If you enter only one "
        "figure, it uses that as the Actual Turnover, which may place the client in the wrong revenue band."
    ),
    (
        "Selecting a cover limit that is too low for the client's turnover band",
        "A client with R80M turnover should not be offered only R2.5M cover. Check the availability matrix "
        "in Section 4.1.1 and ensure you recommend appropriate limits for the revenue band."
    ),
    (
        "Forgetting to switch to 'Renewal' quote type for existing clients",
        "The engine defaults to New Business. If you quote a renewal as New Business, the comparison "
        "will use Industry Benchmark instead of the Existing Policy premium, and the FP equivalent "
        "toggle will not be auto-set. Always confirm the quote type before proceeding."
    ),
    (
        "Not reviewing the audit trail before exporting",
        "The audit trail in Step 5 shows every calculation step. Always verify that the base premium, "
        "modifier, loading, FP cost, and discounts are correct before downloading the PDF."
    ),
    (
        "Using the local engine URL for production quotes",
        "The local URL (localhost:5002) is for development and testing only. Production quotes must use "
        "the live Render URL (https://sme-rating-engine.onrender.com). Quotes generated locally are not "
        "saved to the production database."
    ),
    (
        "Ignoring the 'Refer' flag",
        "Even though the engine allows you to proceed when a 'Refer' flag is active (e.g., Q9 = No), "
        "senior underwriter sign-off is still required. Do not send the quote to the broker without "
        "obtaining this approval."
    ),
]

for i, (title_text, detail) in enumerate(mistakes, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{i}. ")
    run_num.bold = True
    run_num.font.color.rgb = PHISHIELD_RED
    run_title = p.add_run(title_text)
    run_title.bold = True
    run_title.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.75)
    run_detail = p2.add_run(detail)
    run_detail.font.size = Pt(10)
    run_detail.font.color.rgb = PHISHIELD_GREY

add_tip_box(
    "Print this list and keep it next to the Quick Reference Card. A quick check before exporting "
    "can save significant time and avoid compliance issues.",
    "tip"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. TROUBLESHOOTING & FAQ
# ══════════════════════════════════════════════════════════════════════════════

add_heading("9. Troubleshooting & FAQ", level=1)

faqs = [
    (
        "The 'Continue' button is greyed out. What am I missing?",
        "Check the following: (1) Company name is entered, (2) Industry is selected, "
        "(3) At least one turnover figure is entered, (4) Q1 has been answered. "
        "Also check for active blockers: turnover > R200M, Q1 = No (decline), or a referred industry."
    ),
    (
        "I see 'Refer for Underwriting'. Can I still generate a quote?",
        "It depends on the reason. If the industry (Healthcare, Public Administration) triggered the referral, "
        "the engine blocks the quote entirely \u2014 contact a senior underwriter. If Q9 = No triggered it, "
        "you can still proceed through the wizard, but the quote is flagged as 'Refer' and needs senior "
        "underwriter sign-off before being sent to the client."
    ),
    (
        "What is the difference between 'Decline' and 'Refer'?",
        "Decline: The risk does not meet minimum requirements (Q1 = No, or all Q2\u2013Q6 = No). "
        "The quote cannot proceed. Refer: The risk has factors that require senior review (Q9 = No, "
        "prior claim, turnover > R200M). A senior underwriter must assess before proceeding."
    ),
    (
        "Why does my Software & Technology client's premium look higher?",
        "Software & Technology attracts an industry modifier of 1.35x to 1.63x (depending on revenue band) "
        "because tech companies have a higher cyber risk profile. This modifier is applied to the base premium "
        "only, not the Funds Protect cost."
    ),
    (
        "Can I quote the same cover limit with different FP options?",
        "Yes. Use the '+' button on a cover card to add a duplicate. Each duplicate can have a different FP "
        "tier selected. This lets you present the client with options like 'R5M cover with R250k FP' and "
        "'R5M cover with R500k FP' side by side."
    ),
    (
        "What does 'Condition of Cover' mean?",
        "When Q7 or Q8 is answered 'No', the corresponding security measure becomes a condition that the "
        "client must implement. It is noted on the quote output and PDF. The policy may not pay out for "
        "FP-related claims if the client has not implemented the required measures."
    ),
    (
        "My combined discount exceeds 35%. Is this allowed?",
        "The engine allows it but displays a warning. Discounts above 35% require approval from a senior "
        "underwriter. Document the justification before proceeding."
    ),
    (
        "Where are my saved quotes stored?",
        "All quotes are saved to a backend SQLite database automatically. Each quote has a unique reference "
        "number (CPB-YYYYMMDD-NNNN). PDFs are stored in the quote_pdfs directory on the server, organised "
        "by year, month, and company name."
    ),
    (
        "The page is not loading or looks broken.",
        "Try the following: (1) Clear your browser cache, (2) Ensure you are using a supported browser "
        "(Chrome, Edge, or Firefox), (3) Check your internet connection if using the live URL, "
        "(4) Contact IT support if the issue persists."
    ),
    (
        "Can I edit a quote after exporting it?",
        "The engine does not support editing a finalised quote. To change a quote, start a new one with "
        "the updated details. The new quote will receive a new reference number."
    ),
    (
        "The page looks broken or old after an update.",
        "Clear your browser cache. In Chrome: press Ctrl+Shift+Delete, select 'Cached images and files', "
        "then click 'Clear data'. Alternatively, press Ctrl+Shift+R to do a hard refresh which bypasses the cache."
    ),
    (
        "Buttons are not responding when I click them.",
        "Try the following: (1) Hard refresh with Ctrl+Shift+R, (2) Check if a blocker overlay is active "
        "(look for a red or orange message on screen), (3) Try a different browser (Chrome, Edge, or Firefox), "
        "(4) Clear your browser cache completely, (5) If using the live URL, the server may be waking up from "
        "sleep \u2014 wait 30 seconds and try again."
    ),
    (
        "The page takes a long time to load on the live URL.",
        "The live engine runs on Render's free tier, which puts the server to sleep after 15 minutes of "
        "inactivity. The first visit after a sleep period takes 30\u201360 seconds to 'cold start'. Subsequent "
        "visits will be fast. If you use the engine daily, bookmark it and visit it first thing in the morning "
        "to wake it up."
    ),
    (
        "I can't see the latest changes after an update was deployed.",
        "Always hard refresh (Ctrl+Shift+R) after being told an update has been deployed. If that doesn't "
        "work, clear your full browser cache (Ctrl+Shift+Delete \u2192 select all time \u2192 Clear data)."
    ),
    (
        "The PDF download failed or nothing happened when I clicked Download.",
        "(1) Check if your browser is blocking pop-ups or downloads \u2014 look for a small icon in the address "
        "bar, (2) Try disabling any ad blockers for the site, (3) Check your Downloads folder \u2014 the file may "
        "have downloaded but the notification was hidden, (4) Try using Chrome if you're on another browser."
    ),
    (
        "The industry dropdown isn't showing any options.",
        "The industry dropdown is searchable. Start typing the industry name (e.g., 'software', 'retail', "
        "'mining') and options will filter automatically. If nothing appears, clear the search text and try "
        "scrolling through the full list using the mouse wheel."
    ),
]

for q, a in faqs:
    add_bold_para(f"Q: {q}")
    doc.add_paragraph(f"A: {a}")
    doc.add_paragraph()  # spacing

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════

add_heading("10. Quick Reference Card", level=1)

doc.add_paragraph(
    "This page provides a one-page summary of the key rules and thresholds. Print this page "
    "and keep it at your desk for quick reference."
)

doc.add_paragraph()  # spacing

add_bold_para("TURNOVER CALCULATION")
doc.add_paragraph("Actual Turnover = (Previous Year + Current Year) / 2")
doc.add_paragraph("Above R200M = Refer for Underwriting")

doc.add_paragraph()

add_bold_para("REVENUE BANDS")
doc.add_paragraph("R0\u2013R10M | R10M\u2013R25M | R25M\u2013R50M | R50M\u2013R75M | R75M\u2013R100M | R100M\u2013R150M | R150M\u2013R200M")

doc.add_paragraph()

add_bold_para("COVER LIMITS & MINIMUM FP")
make_table(
    ["Cover", "FP Min", "Excess"],
    [
        ["R1M", "R150k", "R10k"],
        ["R2.5M", "R200k", "R10k"],
        ["R5M", "R250k", "R10k"],
        ["R7.5M", "R500k", "R15k"],
        ["R10M", "R1M", "R20k"],
        ["R15M", "R1.5M", "R25k"],
    ]
)

add_bold_para("MICRO SME")
doc.add_paragraph("Turnover < R50M AND Cover <= R5M = Flat-rate Micro SME pricing")

doc.add_paragraph()

add_bold_para("INDUSTRY MODIFIERS (Base Premium Only)")
doc.add_paragraph("Software & Technology: 1.35x \u2013 1.63x")
doc.add_paragraph("Finance/Insurance/Real Estate: 1.28x \u2013 1.46x")
doc.add_paragraph("All others: 1.00x (no modifier)")

doc.add_paragraph()

add_bold_para("UNDERWRITING RULES")
doc.add_paragraph("Q1 = No \u2192 Decline (hard block)")
doc.add_paragraph("Q2\u2013Q6: 0 No = Standard | 1 No = Caution | 2 No = 5% | 3\u20134 No = 10% | 5 No = Decline")
doc.add_paragraph("Q7/Q8 (FP > R250k): No = Condition of Cover")
doc.add_paragraph("Q9 = No \u2192 Refer to senior underwriter")
doc.add_paragraph("Healthcare / Public Admin \u2192 Always Refer")

doc.add_paragraph()

add_bold_para("DISCOUNTS")
doc.add_paragraph("Posture: 0\u201335% | Discretionary: 0\u201335%")
doc.add_paragraph("Combined > 35% = Senior underwriter approval required")
doc.add_paragraph("Applied to full premium (base + FP), compound: Final = Premium x (1-Posture) x (1-Disc)")

doc.add_paragraph()

add_bold_para("QUOTE REFERENCE FORMAT")
doc.add_paragraph("CPB-YYYYMMDD-NNNN-[CoverLimit]-FP[FPLimit]")

doc.add_paragraph()

add_bold_para("QUOTE TYPES")
doc.add_paragraph("New Business | Renewal (shows Existing Policy comparison, FP equiv = Yes) | Competing Quote")

doc.add_paragraph()

add_bold_para("COMPARISON COLOURS")
doc.add_paragraph("Green = Competitive | Amber = Close | Red = Over benchmark")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. GLOSSARY OF KEY TERMS (NEW)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("11. Glossary of Key Terms", level=1)

doc.add_paragraph(
    "This glossary defines the key terms used throughout this manual and in the SME Rating Engine interface."
)

glossary = [
    ("Audit Trail",
     "The step-by-step calculation breakdown showing how the final premium was derived. Displayed in "
     "Step 5 and included in the PDF output."),
    ("Base Premium",
     "The core insurance premium before Funds Protect is added, calculated using formulas based on "
     "turnover and cover limit. For Micro SME, these are flat rates; for standard quotes, they are "
     "formula-based."),
    ("Condition of Cover",
     "A requirement that the client must implement as a condition of their policy. Triggered when "
     "Q7 or Q8 is answered 'No'. Noted on the quote output and PDF."),
    ("Discretionary Discount",
     "A competitive discount applied to match or beat competitor pricing. Entered as a percentage "
     "(0\u201335%) in Step 4 and applied to the full premium."),
    ("Endorsement",
     "Special terms, conditions, or notes added to the quote by the underwriter. Entered in a free-text "
     "area in Step 4 and included on the quote output."),
    ("Funds Protect (FP)",
     "An add-on cover for financial losses from cyber-enabled fraud (business email compromise, social "
     "engineering). Has a minimum limit per cover level, with optional upgrades up to R5M."),
    ("Industry Modifier",
     "A multiplier (1.00x to 1.63x) applied to the base premium for higher-risk industries. Software & "
     "Technology and Finance attract higher modifiers. Applied to base premium only, not FP."),
    ("Micro SME",
     "A flat-rate pricing tier for smaller businesses (turnover < R50M, cover limit R5M or below) "
     "that is lower than standard formula rates. Industry modifiers still apply."),
    ("Posture Discount",
     "A discount based on the quality of the client's cyber security posture, typically supported by "
     "evidence such as SOC 2 certification, penetration testing, or a dedicated CISO."),
    ("Quote Reference",
     "Unique identifier for each quote option in the format CPB-YYYYMMDD-NNNN-CoverLimit-FPLimit. "
     "Multi-cover quotes share the same base reference with unique suffixes."),
    ("Revenue Band",
     "The turnover bracket that determines base premium rates. There are 7 bands from R0 to R200M. "
     "The band is determined by the Actual Turnover (midpoint of previous and current year)."),
    ("Underwriting Loading",
     "A percentage increase (5\u201315%) applied to the base premium when security posture questions "
     "(Q2\u2013Q6) are answered 'No'. The loading percentage depends on the number of 'No' answers."),
]

for term, definition in glossary:
    p = doc.add_paragraph()
    run_term = p.add_run(term)
    run_term.bold = True
    run_term.font.color.rgb = PHISHIELD_ACCENT
    run_term.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.75)
    p2.paragraph_format.space_after = Pt(8)
    run_def = p2.add_run(definition)
    run_def.font.size = Pt(10)
    run_def.font.color.rgb = PHISHIELD_DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. VERSION HISTORY / CHANGELOG (NEW)
# ══════════════════════════════════════════════════════════════════════════════

add_heading("12. Version History / Changelog", level=1)

doc.add_paragraph(
    "This section tracks changes to the SME Rating Engine and this user manual."
)

make_table(
    ["Version", "Date", "Changes"],
    [
        ["1.0", "March 2026",
         "Initial release"],
        ["1.1", "March 2026",
         "Updated S&T modifiers (1.35-1.63x), Finance modifiers (1.28-1.46x), "
         "Micro SME expanded to include S&T and Finance. Added decision flowchart, "
         "glossary, market conditions guide, common mistakes section, expanded worked "
         "examples with detailed Step 3/4 numbers, and version history."],
        ["1.2", "March 2026",
         "Removed Bryte Insurance from cover page. Added IT/Technical FAQ section. "
         "Improved fonts (Calibri 11pt body, 13pt H2, 16pt H1), spacing (1.3 line spacing), "
         "and table readability. Added numbered figure captions. Fixed section numbering."],
        ["1.3", "April 2026",
         "Renewal logic overhaul. Added Current FP Sub-limit as a required renewal field (all three "
         "renewal inputs now mandatory). Q9 auto-set to Yes and locked on Renewal. Prior claim and "
         "Renewal-with-Q9-No contradiction both now hard-block the quote as Refer to Senior UW. "
         "Step 2 recommendations now pin the existing cover as 'Current Cover'. Market-condition "
         "logic extended: softening \u2192 upgrade options, stable \u2192 alternatives, hardening \u2192 "
         "downgrade option. New Premium-drop Protection rule: when new premium at same cover/FP is "
         ">20% below existing, the engine re-recommends a higher cover to retain \u226590% of "
         "existing premium, with a Corporate referral escalation if max SME cover still falls short. "
         "Renewal state fully clears on quote-type switch; step-2 selections re-auto-select when "
         "renewal inputs change. UW loading shown as a per-card badge with a comparison-caveat note. "
         "Section 3.7.3 added; Scenario 8.2b added for the trigger path."],
        ["1.3.1", "April 2026",
         "Rule I ladder-gap refinement. When the Premium-drop target is more than one step above "
         "the existing cover (e.g. R5M \u2192 R10M skipping R7.5M), the engine now fills the skipped "
         "covers as Alternative cards rather than showing only the target plus one cover above it. "
         "Each intermediate Alternative card displays a 'XX% retention' badge showing the new "
         "premium as a percentage of existing at matched FP, so the underwriter can see at a glance "
         "how close each skipped cover sits to the 90% bar. The above-target Alternative is dropped "
         "in the ladder-gap case (no clutter). Adjacent target behaviour (target = existing+1) is "
         "preserved unchanged. Scenario 8.2c added for the ladder-gap worked example."],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER / BACK COVER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()

for _ in range(8):
    doc.add_paragraph()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("End of Document")
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = PHISHIELD_DARK

doc.add_paragraph()

disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run(
    "This document is confidential and intended for internal use by Phishield UMA (Pty) Ltd staff only.\n"
    "Cyber Protect Business Policy (Risk Rated) is underwritten by Bryte Insurance Company Limited.\n\n"
    "For questions or feedback about this manual, contact your team lead or the product team."
)
run.font.size = Pt(9)
run.font.color.rgb = PHISHIELD_GREY

# ── Save ──────────────────────────────────────────────────────────────────────

output_dir = r"C:\Users\sarel\Desktop\Sarel\SML Consulting\PSQ\SME Rating Engine"
docx_path = os.path.join(output_dir, "Phishield SME Rating Engine User Manual.docx")
doc.save(docx_path)
print(f"Manual saved to: {docx_path}")

# ── Generate PDF ──────────────────────────────────────────────────────────────

pdf_path = os.path.join(output_dir, "Phishield SME Rating Engine User Manual.pdf")

try:
    from docx2pdf import convert
    convert(docx_path, pdf_path)
    print(f"PDF saved to: {pdf_path}")
except ImportError:
    print("docx2pdf not installed. Attempting alternative PDF generation...")
    try:
        import subprocess
        # Try using LibreOffice if available
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            print(f"PDF saved to: {pdf_path}")
        else:
            print(f"LibreOffice conversion failed: {result.stderr}")
            print("Install docx2pdf (pip install docx2pdf) or LibreOffice for PDF generation.")
    except FileNotFoundError:
        print("Neither docx2pdf nor LibreOffice found. Install one of them for PDF generation.")
        print("  pip install docx2pdf   (requires Microsoft Word)")
        print("  Or install LibreOffice (free, open-source)")
except Exception as e:
    print(f"PDF generation failed: {e}")
    print("The DOCX file has been saved successfully. Convert it manually if needed.")
