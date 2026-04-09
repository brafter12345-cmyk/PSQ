"""
Phishield Cyber Risk Scanner - User Manual
Part 1: Cover Page, Table of Contents, Section 1 (Introduction), Section 2 (Getting Started)

Usage:
    from part1_intro import build
    build(doc)   # doc is a python-docx Document object
"""

from docx.shared import Emu, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# These constants and helpers are expected to be defined by the caller / main
# script.  We import them at call time via the module-level fallback below so
# that this file can also be tested standalone.

H1_SIZE = Emu(203200)
H2_SIZE = Emu(165100)


# ---------------------------------------------------------------------------
# Local helper stubs -- the main assembler script should supply real versions
# via monkey-patching or by passing them in.  These are thin defaults so the
# file remains independently importable for testing.
# ---------------------------------------------------------------------------

def _fallback_add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = H1_SIZE
    return p


def _fallback_add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = H2_SIZE
    return p


def _fallback_add_body(doc, text):
    return doc.add_paragraph(text)


def _fallback_add_bold_body(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    b = p.add_run(bold_text)
    b.bold = True
    p.add_run(normal_text)
    return p


def _fallback_add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Paragraph")


def _fallback_add_tip(doc, text):
    p = doc.add_paragraph()
    b = p.add_run("TIP: ")
    b.bold = True
    p.add_run(text)
    return p


def _fallback_add_warning(doc, text):
    p = doc.add_paragraph()
    b = p.add_run("WARNING: ")
    b.bold = True
    p.add_run(text)
    return p


def _fallback_add_note(doc, text):
    p = doc.add_paragraph()
    b = p.add_run("NOTE: ")
    b.bold = True
    p.add_run(text)
    return p


# ---------------------------------------------------------------------------
# Helper resolver -- allows the main script to inject real helpers
# ---------------------------------------------------------------------------

_helpers = {}


def set_helpers(**kwargs):
    """Call from main assembler to inject shared helper functions."""
    _helpers.update(kwargs)


def _h(name):
    """Return the helper function, preferring injected over fallback."""
    if name in _helpers:
        return _helpers[name]
    return globals()[f"_fallback_{name}"]


# ---------------------------------------------------------------------------
# Cover Page
# ---------------------------------------------------------------------------

def _build_cover(doc):
    # Add several blank paragraphs for vertical spacing
    for _ in range(6):
        doc.add_paragraph()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PHISHIELD")
    run.bold = True
    run.font.size = Pt(36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cyber Risk Scanner")
    run.bold = True
    run.font.size = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("User Manual")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("External Passive Security Evaluation\nfor Cyber Insurance Underwriting")
    run.font.size = Pt(14)

    for _ in range(4):
        doc.add_paragraph()

    # Metadata block
    meta_lines = [
        "Administrator: Phishield UMA (Pty) Ltd",
        "Version 1.0",
        "April 2026",
        "Confidential",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table of Contents (text-based)
# ---------------------------------------------------------------------------

def _build_toc(doc):
    _h("add_h1")(doc, "Table of Contents")
    doc.add_paragraph()

    toc_entries = [
        ("1", "Introduction"),
        ("  1.1", "Purpose"),
        ("  1.2", "Who Should Use This Manual"),
        ("  1.3", "How the Scanner Works"),
        ("  1.4", "Scanner Architecture"),
        ("  1.5", "Limitations"),
        ("2", "Getting Started"),
        ("  2.1", "Accessing the Scanner"),
        ("  2.2", "Browser Requirements"),
        ("  2.3", "Screen Layout"),
        ("  2.4", "Starting a Scan"),
        ("  2.5", "Scan Progress"),
        ("3", "Understanding Scan Results"),
        ("  3.1", "Risk Score Overview"),
        ("  3.2", "Category Breakdown"),
        ("  3.3", "Checker Detail Cards"),
        ("  3.4", "Recommendations"),
        ("4", "Security Checkers Reference"),
        ("  4.1", "Discovery Checkers"),
        ("  4.2", "Core Security Checkers"),
        ("  4.3", "Information Security Checkers"),
        ("  4.4", "Email Security Checkers"),
        ("  4.5", "Network & Infrastructure Checkers"),
        ("  4.6", "Exposure & Reputation Checkers"),
        ("  4.7", "Technology & Governance Checkers"),
        ("5", "Insurance Analytics"),
        ("  5.1", "Ransomware Susceptibility Index (RSI)"),
        ("  5.2", "FAIR Model & Monte Carlo Simulation"),
        ("  5.3", "Data Breach Index (DBI)"),
        ("  5.4", "Financial Impact Estimates"),
        ("6", "Scoring Methodology"),
        ("  6.1", "Per-Checker Scoring"),
        ("  6.2", "Category Weighting"),
        ("  6.3", "Overall Risk Score Calculation"),
        ("  6.4", "Risk Levels"),
        ("7", "PDF Report"),
        ("  7.1", "Report Structure"),
        ("  7.2", "Executive Summary"),
        ("  7.3", "Technical Detail Pages"),
        ("  7.4", "Downloading & Sharing"),
        ("8", "API Reference"),
        ("  8.1", "POST /api/scan"),
        ("  8.2", "GET /api/scan/<id>"),
        ("  8.3", "GET /api/scan/<id>/pdf"),
        ("  8.4", "GET /api/history/<domain>"),
        ("9", "Configuration & Deployment"),
        ("  9.1", "Environment Variables"),
        ("  9.2", "API Keys"),
        ("  9.3", "Database"),
        ("  9.4", "Deployment on Render"),
        ("10", "Troubleshooting"),
        ("  10.1", "Common Issues"),
        ("  10.2", "Timeout Errors"),
        ("  10.3", "API Rate Limits"),
        ("11", "Glossary"),
    ]

    for number, title in toc_entries:
        indent = number.startswith("  ")
        display_num = number.strip()
        line = f"{display_num}  {title}"
        p = doc.add_paragraph(line)
        if indent:
            p.paragraph_format.left_indent = Inches(0.4)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 1 -- Introduction
# ---------------------------------------------------------------------------

def _build_section_1(doc):
    _h("add_h1")(doc, "1. Introduction")
    doc.add_paragraph()

    # --- 1.1 Purpose ---
    _h("add_h2")(doc, "1.1 Purpose")

    _h("add_body")(doc,
        "The Phishield Cyber Risk Scanner is an external, passive security evaluation tool "
        "designed specifically for cyber insurance underwriting. It examines the publicly "
        "observable digital infrastructure of an organisation -- DNS records, SSL certificates, "
        "HTTP headers, open ports, email authentication, credential leaks, dark-web exposure, "
        "and reputation data -- without ever sending a single intrusive packet to the target."
    )

    _h("add_body")(doc,
        "The scanner was developed by Phishield UMA (Pty) Ltd to support three core insurance "
        "functions:"
    )

    _h("add_bullet")(doc,
        "Risk Assessment -- Produce a quantitative risk profile of a prospective policyholder "
        "before the policy is bound, giving underwriters an objective, data-driven view of the "
        "organisation's external security posture."
    )
    _h("add_bullet")(doc,
        "Premium Calibration -- Feed the scanner's Ransomware Susceptibility Index (RSI), FAIR-model "
        "financial impact estimates, and Data Breach Index (DBI) into the premium "
        "calculation workflow so that pricing reflects actual technical risk rather than "
        "self-reported questionnaire answers alone."
    )
    _h("add_bullet")(doc,
        "Broker Reporting -- Generate a professional PDF report that brokers can present to "
        "clients, explaining findings in plain language alongside actionable recommendations "
        "that can improve the client's risk profile and, in turn, their premium."
    )

    _h("add_body")(doc,
        "Because the scanner operates passively and externally, it can be run on any domain "
        "at any time without the knowledge or cooperation of the target organisation. This "
        "makes it ideal for pre-quote assessments where the prospective policyholder has not "
        "yet engaged technically."
    )

    _h("add_note")(doc,
        "The scanner does not replace a full penetration test or an internal security audit. "
        "It is designed to complement those activities by providing a rapid, low-friction "
        "external view that can be obtained in minutes rather than weeks."
    )

    doc.add_paragraph()

    # --- 1.2 Who Should Use This Manual ---
    _h("add_h2")(doc, "1.2 Who Should Use This Manual")

    _h("add_body")(doc,
        "This manual is intended for three primary audiences:"
    )

    _h("add_bold_body")(doc, "Underwriters ",
        "who need to interpret scan results when assessing cyber insurance applications. "
        "The manual explains what each checker measures, how the overall risk score is "
        "calculated, and what the insurance analytics section (RSI, FAIR, DBI) means in "
        "practical underwriting terms. No deep technical expertise is assumed; the manual "
        "translates technical findings into risk language."
    )
    _h("add_bold_body")(doc, "Brokers ",
        "who present scan reports to prospective or existing clients. The manual describes "
        "the structure of the PDF report, how to read the executive summary and recommendation "
        "sections, and how to explain findings to a non-technical audience. Brokers will also "
        "benefit from the Getting Started section, which explains how to initiate scans and "
        "download reports."
    )
    _h("add_bold_body")(doc, "Technical Staff ",
        "responsible for deploying, configuring, or maintaining the scanner. The manual covers "
        "environment variables, API key management, deployment on Render, the REST API, and "
        "troubleshooting guidance. Technical staff should read the full manual, paying "
        "particular attention to the Configuration & Deployment and API Reference sections."
    )

    _h("add_tip")(doc,
        "If you are a broker seeing the scanner for the first time, start with Section 2 "
        "(Getting Started) and Section 3 (Understanding Scan Results). You can refer to "
        "Section 4 (Security Checkers Reference) when you need to explain a specific finding "
        "to a client."
    )

    doc.add_paragraph()

    # --- 1.3 How the Scanner Works ---
    _h("add_h2")(doc, "1.3 How the Scanner Works")

    _h("add_body")(doc,
        "The Phishield scanner is a passive reconnaissance tool. It collects only information "
        "that is already publicly available on the internet. It does not exploit "
        "vulnerabilities, attempt logins, inject payloads, or modify any data on the target's "
        "systems. In cybersecurity terminology, it performs open-source intelligence (OSINT) "
        "gathering and passive fingerprinting."
    )

    _h("add_body")(doc,
        "When you submit a domain for scanning, the scanner executes a multi-phase pipeline:"
    )

    _h("add_bold_body")(doc, "Phase 1 -- IP Discovery. ",
        "The scanner resolves the domain's DNS A records to obtain all public IP addresses "
        "associated with the domain. If additional client-supplied IP addresses are provided "
        "(for example, known mail server IPs or secondary data-centre addresses), these are "
        "merged into the scan scope. Every subsequent checker operates against either the "
        "domain name itself or the discovered IP addresses."
    )

    _h("add_bold_body")(doc, "Phase 2a -- Lightweight Domain Checkers (Concurrent). ",
        "Approximately 20 checkers run simultaneously against the domain. These include "
        "email authentication (SPF, DKIM, DMARC), HTTP header analysis, WAF detection, "
        "technology stack fingerprinting, breach database lookups, credential leak checks, "
        "dark-web exposure searches, admin panel detection, domain intelligence, privacy "
        "compliance analysis, and more. Running these concurrently keeps total scan time "
        "under control."
    )

    _h("add_bold_body")(doc, "Phase 2b -- Heavy Checkers (Sequential). ",
        "A small number of resource-intensive checkers run one at a time after the concurrent "
        "batch completes. These include the SSL/TLS deep analysis (which spawns a subprocess "
        "to evaluate cipher suites, certificate chains, and protocol support) and subdomain "
        "reconnaissance (which queries Certificate Transparency logs and resolves large "
        "numbers of hostnames). Running these sequentially prevents memory exhaustion on "
        "resource-constrained hosting environments."
    )

    _h("add_bold_body")(doc, "Phase 3 -- IP-Level Checkers. ",
        "For each discovered IP address, the scanner runs a set of IP-specific checkers: "
        "DNS infrastructure and open port scanning, high-risk protocol detection (RDP, "
        "Telnet, FTP, SMB), Shodan vulnerability lookups, and DNSBL blacklist checks. "
        "Results from multiple IPs are merged into a single aggregate view, preserving "
        "per-IP detail for the technical report."
    )

    _h("add_bold_body")(doc, "Phase 4 -- OSV Enrichment. ",
        "Any technologies or software versions detected during earlier phases are cross-"
        "referenced against the OSV.dev vulnerability database (which includes the National "
        "Vulnerability Database / NVD). This enrichment step attaches known CVEs to detected "
        "software, enabling the report to highlight specific, exploitable vulnerabilities "
        "rather than generic warnings."
    )

    _h("add_bold_body")(doc, "Phase 5 -- Scoring. ",
        "Each checker produces a normalised score. These are combined using category-level "
        "weighting to produce an overall risk score from 0 (critical risk) to 100 (excellent "
        "posture). The risk level is then classified as Critical, High, Medium, Low, or "
        "Excellent based on defined thresholds."
    )

    _h("add_bold_body")(doc, "Phase 6 -- Insurance Analytics. ",
        "The scanner calculates three insurance-specific metrics: the Ransomware Susceptibility Index "
        "(RSI), which adjusts the technical score for industry risk and organisational size; "
        "FAIR-model financial impact estimates using Monte Carlo simulation with 10,000 "
        "iterations; and the Data Breach Index (DBI), which estimates the organisation's "
        "digital dependency. These analytics feed directly into premium calculation and "
        "underwriting decision-making."
    )

    _h("add_note")(doc,
        "The entire scan typically completes in 5 to 15 minutes depending on the complexity "
        "of the target domain and the responsiveness of external APIs. Progress is streamed "
        "to the browser in real time via Server-Sent Events (SSE)."
    )

    doc.add_paragraph()

    # --- 1.4 Scanner Architecture ---
    _h("add_h2")(doc, "1.4 Scanner Architecture")

    _h("add_body")(doc,
        "The scanner comprises over 25 individual checkers organised into eight categories. "
        "Each category groups related security concerns together, making it easier for "
        "underwriters and brokers to quickly identify where the target's strengths and "
        "weaknesses lie."
    )

    _h("add_bold_body")(doc, "Discovery (2 checkers): ",
        "IP Discovery and Web Ranking. Establishes the scan scope by resolving all public "
        "IP addresses and assessing the domain's web traffic ranking to provide context on "
        "organisational visibility."
    )
    _h("add_bold_body")(doc, "Core Security (4 checkers): ",
        "SSL/TLS Certificate, HTTP Security Headers, Website Security, and WAF/DDoS "
        "Protection. These evaluate the fundamental security controls that every "
        "internet-facing organisation should have in place."
    )
    _h("add_bold_body")(doc, "Information Security (1 checker): ",
        "Information Disclosure. Detects unintentional exposure of sensitive technical "
        "details such as server software versions, directory listings, debug pages, "
        "and configuration files."
    )
    _h("add_bold_body")(doc, "Email Security (2 checkers): ",
        "Email Authentication (SPF, DKIM, DMARC) and Email Hardening (MTA-STS, "
        "TLS-RPT, DANE). These assess the organisation's resilience to email spoofing, "
        "phishing, and business email compromise (BEC) attacks."
    )
    _h("add_bold_body")(doc, "Network & Infrastructure (6 checkers): ",
        "DNS & Open Ports, High-Risk Protocols, Shodan Vulnerabilities, DNSBL/Blacklists, "
        "Cloud & CDN, and VPN/Remote Access. These checkers evaluate the network perimeter, "
        "identifying open ports, exposed services, known vulnerabilities, and blacklist "
        "appearances. Several of these run per-IP to cover all resolved addresses."
    )
    _h("add_bold_body")(doc, "Exposure & Reputation (6 checkers): ",
        "Data Breaches (HIBP), Credential Leaks (Dehashed), Exposed Admin Panels, "
        "VirusTotal Intelligence, Subdomain Reconnaissance, and Lookalike/Fraudulent "
        "Domains. These measure the organisation's exposure in breach databases, dark-web "
        "markets, and typosquatting registrations."
    )
    _h("add_bold_body")(doc, "Technology & Governance (6 checkers): ",
        "Technology Stack, Domain Intelligence, SecurityTrails DNS History, Security "
        "Policy & Vulnerability Disclosure Programme (VDP), Payment Security (PCI-DSS "
        "indicators), and Privacy Compliance (cookie consent, privacy policy). These "
        "assess the maturity of the organisation's technology governance and regulatory "
        "compliance posture."
    )
    _h("add_bold_body")(doc, "Insurance Analytics (1 checker): ",
        "RSI / Financial Impact / DBI. This final stage synthesises all preceding results "
        "into insurance-relevant metrics, including annualised loss expectancy calculated "
        "via FAIR-model Monte Carlo simulation."
    )

    _h("add_body")(doc,
        "The scanner integrates with 10 external APIs and data sources to enrich its "
        "findings:"
    )

    _h("add_bullet")(doc, "Shodan -- Internet-wide port and service scanning data, vulnerability lookups")
    _h("add_bullet")(doc, "Have I Been Pwned (HIBP) -- Breach database for domain-level email exposure")
    _h("add_bullet")(doc, "Dehashed -- Credential leak database with password hash and plaintext detection")
    _h("add_bullet")(doc, "VirusTotal -- Malware, phishing, and reputation intelligence")
    _h("add_bullet")(doc, "SecurityTrails -- Historical DNS records, domain ownership changes")
    _h("add_bullet")(doc, "IntelX (Intelligence X) -- Dark-web and paste-site exposure monitoring")
    _h("add_bullet")(doc, "OSV.dev / NVD -- Open-source vulnerability database for CVE enrichment")
    _h("add_bullet")(doc, "HudsonRock -- Infostealer malware credential compromise data")
    _h("add_bullet")(doc, "Certificate Transparency Logs -- Subdomain enumeration via issued certificates")
    _h("add_bullet")(doc, "Multiple DNSBL providers -- IP and domain blacklist status checks")

    _h("add_body")(doc,
        "The scanner produces both an interactive web-based results dashboard and a "
        "comprehensive PDF report. The PDF includes an executive summary, per-category "
        "detail pages, a full recommendations list with prioritisation, and the insurance "
        "analytics section with FAIR-model outputs."
    )

    _h("add_tip")(doc,
        "For a visual overview of which checkers run in each phase and how data flows "
        "through the pipeline, refer to the architecture diagram in the Phishield technical "
        "documentation pack."
    )

    doc.add_paragraph()

    # --- 1.5 Limitations ---
    _h("add_h2")(doc, "1.5 Limitations")

    _h("add_body")(doc,
        "While the Phishield scanner provides a comprehensive external view of an "
        "organisation's security posture, it is important to understand what it cannot "
        "assess. Being transparent about these limitations ensures that underwriters and "
        "brokers set appropriate expectations when using scan results in their decision-making."
    )

    _h("add_bold_body")(doc, "External-only perspective. ",
        "The scanner can only observe what is visible from the public internet. It cannot "
        "assess internal network segmentation, endpoint protection, backup strategies, "
        "staff security awareness training, incident response procedures, or physical "
        "security controls. An organisation could score well externally while having "
        "significant internal weaknesses. The scanner should be used alongside "
        "questionnaires and, where appropriate, internal audits."
    )

    _h("add_bold_body")(doc, "Point-in-time assessment. ",
        "Each scan captures the state of the target's infrastructure at the moment the "
        "scan is executed. Configurations change, new vulnerabilities are disclosed, and "
        "certificates expire. A scan that shows a clean posture today may not reflect "
        "the state a month from now. For ongoing risk monitoring, periodic re-scanning "
        "is recommended. Continuous monitoring capabilities are planned for a future "
        "release."
    )

    _h("add_bold_body")(doc, "External API dependency. ",
        "Several checkers rely on third-party APIs (Shodan, HIBP, Dehashed, VirusTotal, "
        "SecurityTrails, IntelX, HudsonRock). If any of these APIs are unavailable, rate-"
        "limited, or experiencing outages, the corresponding checker may return incomplete "
        "results or an error status. The scanner handles these gracefully by reporting the "
        "error without crashing, but the affected data will be missing from the report."
    )

    _h("add_bold_body")(doc, "HTTPS and network timeouts. ",
        "Some target domains may have aggressive firewall rules, geo-blocking, or very "
        "slow response times that cause individual checkers to time out. When this occurs, "
        "the checker reports a timeout error rather than incomplete data. The overall scan "
        "still completes, but the affected checker's section will show limited or no "
        "findings."
    )

    _h("add_bold_body")(doc, "Not a penetration test. ",
        "The scanner does not attempt to exploit any vulnerabilities it discovers. It does "
        "not test for SQL injection, cross-site scripting, authentication bypass, or any "
        "other active attack vector. The presence of an open port or an outdated software "
        "version is flagged as a risk indicator, but the scanner does not verify whether "
        "that risk is actually exploitable. For exploitability validation, a separate "
        "penetration test is required."
    )

    _h("add_bold_body")(doc, "Subdomain and IP coverage. ",
        "The scanner focuses on the primary domain and its DNS-resolved IP addresses. "
        "While subdomain reconnaissance via Certificate Transparency logs provides broad "
        "coverage, it may not discover every subdomain (for example, those that have never "
        "had a certificate issued). Similarly, IP addresses not linked via DNS A records "
        "or client-supplied lists will not be scanned."
    )

    _h("add_warning")(doc,
        "Never rely on the scanner as the sole basis for an underwriting decision. It is "
        "one input among several and should be considered alongside application forms, "
        "claims history, industry benchmarks, and where available, internal security "
        "assessments."
    )

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 2 -- Getting Started
# ---------------------------------------------------------------------------

def _build_section_2(doc):
    _h("add_h1")(doc, "2. Getting Started")
    doc.add_paragraph()

    # --- 2.1 Accessing the Scanner ---
    _h("add_h2")(doc, "2.1 Accessing the Scanner")

    _h("add_body")(doc,
        "The Phishield Cyber Risk Scanner is a web application that runs in your browser. "
        "No software installation is required. There are two ways to access it:"
    )

    _h("add_bold_body")(doc, "Local Development Instance: ",
        "http://localhost:5001. This is used by technical staff running the scanner on "
        "their own machine for development, testing, or offline use. The local instance "
        "requires Python 3.10+ and the scanner's dependencies to be installed."
    )

    _h("add_bold_body")(doc, "Production Instance: ",
        "https://phishield-scanner.onrender.com. This is the hosted version used by "
        "underwriters and brokers for day-to-day scanning. It is deployed on Render and "
        "is accessible from any internet-connected device."
    )

    _h("add_note")(doc,
        "The production instance may take 30-60 seconds to respond on the first request "
        "if it has been idle, because the Render free tier spins down inactive services. "
        "Subsequent requests will respond immediately."
    )

    _h("add_tip")(doc,
        "Bookmark the production URL for quick access. If you experience slow loading, "
        "simply wait for the initial spin-up to complete and then refresh the page."
    )

    doc.add_paragraph()

    # --- 2.2 Browser Requirements ---
    _h("add_h2")(doc, "2.2 Browser Requirements")

    _h("add_body")(doc,
        "The scanner's web interface is built with standard HTML, CSS, and JavaScript and "
        "is compatible with all modern browsers. For the best experience, use one of the "
        "following:"
    )

    _h("add_bullet")(doc, "Google Chrome (version 90 or later) -- recommended")
    _h("add_bullet")(doc, "Microsoft Edge (version 90 or later)")
    _h("add_bullet")(doc, "Mozilla Firefox (version 88 or later)")
    _h("add_bullet")(doc, "Apple Safari (version 14 or later)")

    _h("add_body")(doc,
        "The scanner uses Server-Sent Events (SSE) to stream real-time progress updates "
        "to your browser during a scan. All of the browsers listed above support SSE "
        "natively. Internet Explorer is not supported."
    )

    _h("add_warning")(doc,
        "Do not use browser extensions that aggressively block JavaScript or network "
        "requests (such as NoScript or certain ad-blockers in strict mode), as they may "
        "interfere with the scan progress stream and results display."
    )

    doc.add_paragraph()

    # --- 2.3 Screen Layout ---
    _h("add_h2")(doc, "2.3 Screen Layout")

    _h("add_body")(doc,
        "The scanner interface is divided into three main areas that you will interact "
        "with during a typical scan workflow:"
    )

    _h("add_bold_body")(doc, "1. Scan Form (Top of Page). ",
        "This is where you enter the target domain and configure scan options. The form "
        "contains the domain input field, industry dropdown, annual revenue field, and "
        "several toggle switches for optional data sources. A prominent 'Start Scan' "
        "button initiates the scan. This area is visible when no scan is in progress."
    )

    _h("add_bold_body")(doc, "2. Progress View (During Scan). ",
        "Once a scan is started, the form is replaced by a real-time progress dashboard. "
        "This shows an elapsed time counter, the domain being scanned, and a grid of "
        "all checkers grouped by category. Each checker displays its current status: "
        "a spinning indicator while running, a green tick when completed successfully, "
        "or a red cross if an error occurred. The checker's score or grade is shown "
        "next to its name as soon as it finishes."
    )

    _h("add_bold_body")(doc, "3. Results Dashboard (After Scan). ",
        "When all checkers have completed, the progress view transitions to the results "
        "dashboard. This displays the overall risk score prominently at the top, followed "
        "by a category-by-category breakdown. Each category can be expanded to reveal "
        "individual checker results, which in turn can be expanded further to show "
        "detailed findings, issues, and recommendations. A 'Download PDF' button is "
        "available to generate and download the full report."
    )

    _h("add_tip")(doc,
        "You can leave the browser tab open during a scan and switch to other tasks. "
        "The scan runs server-side, so closing the progress view does not cancel the scan. "
        "You can return to view results using the scan's unique URL."
    )

    doc.add_paragraph()

    # --- 2.4 Starting a Scan ---
    _h("add_h2")(doc, "2.4 Starting a Scan")

    _h("add_body")(doc,
        "To initiate a new scan, navigate to the scanner's home page and fill in the "
        "scan form. Each field is described below:"
    )

    _h("add_bold_body")(doc, "Domain (required). ",
        "Enter the primary domain you wish to scan, for example: example.co.za. Do not "
        "include the protocol prefix (https:// or http://) or any path. The scanner will "
        "automatically strip these if entered. Only one domain can be scanned at a time. "
        "Subdomains (such as www.example.co.za) can be entered but the scanner will focus "
        "on the root domain for most checks."
    )

    _h("add_bold_body")(doc, "Industry (dropdown). ",
        "Select the industry that best describes the target organisation. The industry "
        "selection affects two important calculations: the Ransomware Susceptibility Index (RSI) "
        "multiplier, which adjusts the risk score based on the threat landscape for that "
        "industry (financial services and healthcare face higher threat levels than "
        "agriculture, for example); and the FAIR model's threat event frequency and loss "
        "magnitude parameters, which vary significantly by sector. Available industries "
        "include Agriculture, Communications, Consumer, Education, Energy, Entertainment, "
        "Financial Services, Healthcare, Hospitality, Industrial/Manufacturing, Legal, "
        "Media, Pharmaceuticals, Public Sector, Research, Retail, Services, Technology, "
        "Transportation, Government, and Other."
    )

    _h("add_bold_body")(doc, "Annual Revenue (ZAR). ",
        "Enter the target organisation's annual revenue in South African Rand. This value "
        "is used to calculate the size multiplier for the RSI (larger organisations "
        "represent larger risk exposure) and to calibrate the FAIR model's financial "
        "impact estimates (loss magnitude scales with revenue). If the exact figure is "
        "not known, use a reasonable estimate. Entering zero will cause the financial "
        "impact estimates to default to industry averages."
    )

    _h("add_bold_body")(doc, "Include Dehashed (toggle). ",
        "When enabled, the scanner queries the Dehashed credential leak database to check "
        "whether email addresses associated with the target domain appear in known data "
        "breaches with exposed passwords. This provides deeper insight than HIBP alone, "
        "as Dehashed includes password hashes and sometimes plaintext credentials. "
        "Enabling this toggle uses paid API credits."
    )

    _h("add_bold_body")(doc, "Include IntelX (toggle). ",
        "When enabled, the scanner queries the Intelligence X (IntelX) database for "
        "dark-web and paste-site mentions of the target domain. This can reveal leaked "
        "documents, credentials posted on underground forums, or other sensitive data "
        "exposures. Enabling this toggle uses paid API credits."
    )

    _h("add_bold_body")(doc, "Include Fraudulent Domains (toggle). ",
        "When enabled, the scanner checks for lookalike (typosquatted) domain registrations "
        "that could be used for phishing attacks against the target organisation's customers "
        "or employees. This checker generates permutations of the target domain and checks "
        "which ones are registered. Enabling this adds approximately 30 seconds to the "
        "total scan time."
    )

    _h("add_note")(doc,
        "The Dehashed and IntelX toggles are off by default to conserve paid API credits. "
        "Enable them when you need the most comprehensive scan possible, such as for "
        "high-value underwriting decisions or when a client has specifically requested "
        "dark-web monitoring."
    )

    _h("add_warning")(doc,
        "Ensure you have authorisation to scan the target domain. While the scanner is "
        "passive and non-intrusive, organisational policies or client agreements may "
        "require prior consent before conducting any form of security assessment."
    )

    _h("add_body")(doc,
        "Once all fields are completed, click the 'Start Scan' button. The scanner will "
        "validate the domain, assign a unique scan identifier, and begin the multi-phase "
        "scanning pipeline. You will be redirected to the progress view automatically."
    )

    doc.add_paragraph()

    # --- 2.5 Scan Progress ---
    _h("add_h2")(doc, "2.5 Scan Progress")

    _h("add_body")(doc,
        "After starting a scan, the browser displays a real-time progress view that keeps "
        "you informed about which checkers are running and which have completed."
    )

    _h("add_bold_body")(doc, "Elapsed Time Counter. ",
        "A timer at the top of the progress view shows how long the scan has been running. "
        "A typical scan completes in 5 to 15 minutes. Complex domains with many subdomains, "
        "multiple IP addresses, or slow-responding servers may take longer."
    )

    _h("add_bold_body")(doc, "Checker Status Grid. ",
        "All checkers are displayed in a grid, grouped by their category (Discovery, Core "
        "Security, Email Security, and so on). Each checker shows one of three states:"
    )

    _h("add_bullet")(doc,
        "Running -- A spinning indicator shows that the checker is currently executing. "
        "Multiple checkers run concurrently during Phase 2a."
    )
    _h("add_bullet")(doc,
        "Done -- A green tick indicates successful completion. The checker's score or "
        "grade appears next to its name."
    )
    _h("add_bullet")(doc,
        "Error -- A red cross indicates that the checker encountered an error, typically "
        "a timeout or API failure. The affected data will be missing from the results."
    )

    _h("add_bold_body")(doc, "Server-Sent Events (SSE). ",
        "Progress updates are delivered to your browser via Server-Sent Events, a "
        "lightweight streaming protocol. This means you see updates in real time without "
        "needing to refresh the page. If your network connection is interrupted briefly, "
        "the browser will automatically reconnect and resume receiving updates."
    )

    _h("add_bold_body")(doc, "Scan Completion. ",
        "When all checkers have finished, the insurance analytics module runs its "
        "calculations (RSI, FAIR Monte Carlo, DBI), the overall risk score is computed, "
        "and the progress view automatically transitions to the full results dashboard. "
        "The scan results are stored in the database and can be accessed at any time "
        "using the scan's unique URL."
    )

    _h("add_tip")(doc,
        "If you need to scan the same domain again later (for example, after the client "
        "has remediated findings), simply start a new scan. The scanner maintains a history "
        "of all scans for each domain, allowing you to compare results over time."
    )

    _h("add_note")(doc,
        "If a scan appears stuck for more than 20 minutes, it is likely that one or more "
        "external APIs are experiencing issues. You can check the scan status via the API "
        "endpoint GET /api/scan/<id> or simply start a new scan. Stalled scans do not "
        "consume ongoing resources."
    )

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def build(doc):
    """Add Cover Page, TOC, Section 1, and Section 2 to *doc*."""
    _build_cover(doc)
    _build_toc(doc)
    _build_section_1(doc)
    _build_section_2(doc)
