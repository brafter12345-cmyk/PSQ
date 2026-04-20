"""Shared helpers for manual part modules — consistent formatting."""
from docx.shared import Emu, Pt, RGBColor

H1_SIZE = Emu(203200)
H2_SIZE = Emu(165100)
H3_SIZE = Pt(11)
NAVY = RGBColor(0, 51, 102)


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = H1_SIZE
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = H2_SIZE
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = H3_SIZE
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.widow_control = True
    return p


def add_bold_body(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(normal_text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.widow_control = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Paragraph')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.widow_control = True
    return p


def add_tip(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("TIP: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    return p


def add_warning(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("WARNING: ")
    r.bold = True
    r.font.color.rgb = RGBColor(180, 0, 0)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.widow_control = True
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("NOTE: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.widow_control = True
    return p
