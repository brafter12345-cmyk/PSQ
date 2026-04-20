"""
Generate Phishield Cyber Risk Scanner User Manual (.docx)
Run: py generate_manual.py
"""

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "Phishield_Cyber_Risk_Scanner_User_Manual.docx")

doc = Document()

# -- Default font --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)

# Helper functions — consistent spacing, keep_with_next to prevent orphaning
def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Emu(203200)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Emu(165100)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    return p

def add_h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.widow_control = True
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Paragraph")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.widow_control = True
    return p

def add_callout(label, text):
    """Add a TIP / WARNING / NOTE callout."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    if label == "WARNING":
        r.font.color.rgb = RGBColor(180, 0, 0)
    elif label == "TIP":
        r.font.color.rgb = RGBColor(0, 128, 0)
    else:
        r.font.color.rgb = RGBColor(0, 51, 153)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return t

def add_page_break():
    doc.add_page_break()


# ============================================================================
# COVER PAGE
# ============================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PHISHIELD")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cyber Risk Scanner")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("User Manual")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("External Passive Security Evaluation for Cyber Insurance Underwriting").font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

for line in [
    "Administrator: Phishield UMA (Pty) Ltd",
    "Version 1.0 | April 2026",
    "",
    "Confidential \u2014 Internal Use Only",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(10)
    if "Confidential" in line:
        r.bold = True
        r.font.color.rgb = RGBColor(180, 0, 0)

add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_h1("Table of Contents")
toc_items = [
    "1.  Introduction",
    "    1.1  Purpose",
    "    1.2  Who Should Use This Manual",
    "    1.3  How the Scanner Works (Overview)",
    "    1.4  Scanner Architecture",
    "    1.5  Limitations",
    "2.  Getting Started",
    "    2.1  Accessing the Scanner",
    "    2.2  Browser Requirements",
    "    2.3  Screen Layout",
    "    2.4  Starting a Scan",
    "3.  Understanding the Results Dashboard",
    "    3.1  Risk Score (0\u20131000)",
    "    3.2  Executive Summary Cards",
    "    3.3  Vulnerability Posture",
    "    3.4  Attacker\u2019s View (Kill Chain)",
    "    3.5  Insurance Analytics",
    "4.  Detailed Checker Explanations",
    "    4.1  Discovery",
    "    4.2  Core Security",
    "    4.3  Information Security",
    "    4.4  Email Security",
    "    4.5  Network & Infrastructure",
    "    4.6  Exposure & Reputation",
    "    4.7  Technology & Governance",
    "    4.8  Compliance Framework Mapping",
    "5.  Insurance Analytics (Deep Dive)",
    "    5.1  Ransomware Susceptibility Index (RSI)",
    "    5.2  Data Breach Index (DBI)",
    "    5.3  Financial Impact (Hybrid Model)",
    "    5.4  Remediation Roadmap",
    "6.  PDF Reports",
    "    6.1  Full Technical Report",
    "    6.2  Broker Summary",
    "    6.3  Reading the Report",
    "7.  Scoring Methodology",
    "    7.1  Overall Risk Score",
    "    7.2  Scoring Failsafe",
    "    7.3  Scan Completeness",
    "    7.4  Score Reliability",
    "8.  API Integrations & Data Sources",
    "9.  Known Limitations & Planned Improvements",
    "10. Glossary",
    "11. Version History",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(10)

add_page_break()

# ============================================================================
# 1. INTRODUCTION
# ============================================================================
add_h1("1. Introduction")

add_h2("1.1 Purpose")
add_body(
    "The Phishield Cyber Risk Scanner is an external, passive security assessment tool purpose-built for "
    "cyber insurance underwriting. It evaluates the publicly visible security posture of any organisation "
    "by scanning its internet-facing infrastructure \u2014 websites, email systems, DNS records, IP addresses, "
    "and exposed services \u2014 without ever attempting to intrude, authenticate, or exploit any system."
)
add_body(
    "The scanner produces a quantitative risk score (0\u20131000), a set of insurance-specific analytics "
    "(Ransomware Susceptibility Index, Data Breach Index, and hybrid financial impact projections), "
    "and a prioritised remediation roadmap. These outputs are designed to support:"
)
add_bullet("Underwriting decisions \u2014 objective, data-driven risk assessment for cyber policy pricing.")
add_bullet("Broker presentations \u2014 clear, visual reports that explain cyber risk to non-technical clients.")
add_bullet("Pre-bind risk improvement \u2014 actionable remediation steps that reduce risk before policy inception.")
add_bullet("Renewal benchmarking \u2014 comparative scans showing security posture changes over time.")

add_h2("1.2 Who Should Use This Manual")
add_body("This manual is intended for the following audiences:")
add_bullet("Underwriters \u2014 to interpret scan results and insurance analytics when pricing cyber policies.")
add_bullet("Brokers and Intermediaries \u2014 to understand the reports they present to prospective policyholders.")
add_bullet("Technical Staff \u2014 to understand the scanning methodology, API integrations, and scoring model in depth.")
add_bullet("Administrators \u2014 to manage the scanner deployment, API keys, and configuration.")

add_callout("NOTE", "No programming or cybersecurity expertise is required to use the scanner or read its reports. "
            "Technical terms are explained in context and defined in the Glossary (Section 10).")

add_h2("1.3 How the Scanner Works (Overview)")
add_body(
    "The scanner operates entirely from the outside, mimicking the reconnaissance an attacker would perform "
    "before targeting an organisation. It collects only publicly available information \u2014 the same data any "
    "person on the internet can access. This means:"
)
add_bullet("No software is installed on the target\u2019s systems.")
add_bullet("No credentials are required from the target organisation.")
add_bullet("No vulnerability exploitation is attempted.")
add_bullet("No traffic is generated that would disrupt normal operations.")
add_body(
    "The scanning process follows a structured multi-phase approach:"
)
add_bullet("Phase 1: IP Discovery \u2014 resolve the domain\u2019s DNS records to find all IP addresses, including any client-supplied IPs.")
add_bullet("Phase 2: Domain-level checks \u2014 21 checkers run concurrently against the domain (SSL, email, headers, WAF, breaches, etc.).")
add_bullet("Phase 3: IP-level checks \u2014 4 checkers run per discovered IP (DNS/ports, high-risk protocols, blacklists, Shodan CVEs).")
add_bullet("Phase 4: Aggregation and enrichment \u2014 merge per-IP results, OSV.dev CVE enrichment, credential risk classification, EPSS and CISA KEV lookups.")
add_bullet("Phase 5: Scoring \u2014 calculate weighted risk score (0\u20131000) across all 25 scoring categories.")
add_bullet("Phase 6: Insurance analytics \u2014 RSI, DBI, hybrid financial impact (Monte Carlo simulation), and remediation roadmap.")

add_callout("TIP", "A typical scan takes 10\u201318 minutes depending on the target\u2019s infrastructure complexity and API response times. "
            "Progress is shown in real time via server-sent events.")

add_h2("1.4 Scanner Architecture")
add_body(
    "The scanner comprises 25+ individual checkers organised into 8 categories, backed by 10 external API "
    "integrations. The architecture is designed for deployment on resource-constrained environments (e.g. Render "
    "free tier with 512 MB RAM) by splitting checkers into lightweight concurrent batches and heavyweight "
    "sequential tasks."
)
add_body("Checker categories and counts:")
add_table(
    ["Category", "Checkers", "Count"],
    [
        ["Discovery", "IP Discovery, Web Ranking", "2"],
        ["Core Security", "SSL/TLS, HTTP Headers, Website Security, WAF", "4"],
        ["Information Security", "Information Disclosure", "1"],
        ["Email Security", "Email Authentication (SPF/DKIM/DMARC), Email Hardening (MTA-STS/TLS-RPT/DANE/BIMI)", "2"],
        ["Network & Infrastructure", "DNS & Open Ports, High-Risk Protocols, Shodan CVEs, DNSBL, Cloud/CDN, VPN/Remote Access", "6"],
        ["Exposure & Reputation", "HIBP Breaches, Dehashed, Exposed Admin, VirusTotal, Subdomains, Hudson Rock, IntelX, Credential Risk, Fraudulent Domains", "9"],
        ["Technology & Governance", "Tech Stack/EOL, Domain Intel, SecurityTrails, Security Policy, Payment Security, Privacy Compliance", "6"],
        ["Insurance Analytics", "RSI, DBI, Financial Impact (Hybrid), Remediation Roadmap", "4"],
    ]
)

add_h2("1.5 Limitations")
add_body("Users should be aware of the following inherent limitations of external-only scanning:")
add_bullet("External only \u2014 the scanner cannot see internal networks, endpoint protection, backup policies, MFA configuration, or staff security awareness. These factors significantly affect actual risk but are invisible from the outside.")
add_bullet("Point-in-time \u2014 results reflect the target\u2019s posture at the moment of scanning. Security configurations can change at any time.")
add_bullet("Not a penetration test \u2014 the scanner does not attempt exploitation. A clean scan does not mean the organisation is invulnerable.")
add_bullet("API dependency \u2014 several checkers rely on third-party APIs (Shodan, VirusTotal, HIBP, etc.) whose availability, rate limits, and data freshness vary.")
add_bullet("Rate limiting \u2014 target websites may rate-limit or block scanner requests, resulting in incomplete data for some checks.")
add_bullet("CDN/WAF masking \u2014 organisations behind CDNs or WAFs may have their true infrastructure hidden, which can make the scan appear more favourable than reality.")

add_callout("WARNING", "The scanner is a risk assessment tool, not a guarantee of security. It should be used as one input "
            "among several in the underwriting process, alongside proposal forms, claims history, and interviews.")

add_page_break()

# ============================================================================
# 2. GETTING STARTED
# ============================================================================
add_h1("2. Getting Started")

add_h2("2.1 Accessing the Scanner")
add_body("The scanner is accessible via two deployment environments:")
add_bullet("Production (Render): https://security-scanner-XXXX.onrender.com \u2014 the live deployment used for client scans.")
add_bullet("Local development: http://localhost:5000 \u2014 for testing and development purposes.")
add_body(
    "Access is currently restricted to authorised Phishield staff. There is no public-facing registration. "
    "Contact your administrator for access credentials if required."
)

add_h2("2.2 Browser Requirements")
add_body("The scanner\u2019s web interface is compatible with all modern browsers:")
add_bullet("Google Chrome 90+ (recommended)")
add_bullet("Mozilla Firefox 88+")
add_bullet("Microsoft Edge 90+")
add_bullet("Safari 14+")
add_callout("NOTE", "JavaScript must be enabled. The results dashboard uses server-sent events (SSE) for real-time "
            "progress updates, which require a persistent HTTP connection.")

add_h2("2.3 Screen Layout")
add_body("The scanner interface consists of three main views:")
add_bullet("Scan Form \u2014 where you enter the domain and scan parameters (industry, revenue, optional toggles).")
add_bullet("Progress View \u2014 a real-time checklist showing each checker as it runs, completes, or fails. Results stream in progressively via SSE.")
add_bullet("Results Dashboard \u2014 the full interactive report with risk score, executive summary cards, vulnerability posture, attacker\u2019s view, insurance analytics, and detailed checker results.")
add_body(
    "The results dashboard is also accessible via a permanent URL: /results/<scan_id>. "
    "This allows you to revisit any previous scan result without re-scanning."
)

add_h2("2.4 Starting a Scan")
add_body("To initiate a scan, provide the following information:")

add_table(
    ["Field", "Required", "Description"],
    [
        ["Domain", "Yes", "The primary domain to scan (e.g. example.co.za). Do not include https:// or paths."],
        ["Industry", "Recommended", "Select the closest matching industry. Used for RSI multipliers and financial impact calculations. Options include: Agriculture, Communications, Consumer, Education, Energy, Entertainment, Financial Services, Healthcare, Hospitality, Industrial/Manufacturing, Legal, Media, Pharmaceuticals, Public Sector, Research, Retail, Services, Technology, Transportation, Government, Other."],
        ["Sub-Industry", "Recommended", "Appears after selecting an industry. Lists sub-industries for the selected sector. Refines the financial impact model for the specific sector. Type to search/filter."],
        ["Annual Revenue (ZAR)", "Recommended", "Annual revenue in South African Rand. Used for financial impact projections. If omitted, a default of R10,000,000 is assumed."],
        ["GDPR Applicable", "Optional", "Enable if the company processes EU personal data. Adds GDPR regulatory exposure (4% of global turnover, uncapped) to the model."],
        ["PCI DSS Applicable", "Optional", "Enable if the company stores/processes card data. Adds PCI fine exposure to the model."],
        ["Include Dehashed", "Optional", "Toggle on to query the Dehashed API for leaked credentials. Requires valid API key."],
        ["Include IntelX", "Optional", "Toggle on to query Intelligence X for dark web exposure. Requires valid API key."],
        ["Include Fraudulent Domains", "Optional", "Toggle on to scan for lookalike/typosquatting domains. This is a heavyweight check and adds several minutes."],
        ["Client IPs", "Optional", "Additional IP addresses to scan (e.g. office IPs, secondary hosting). These are merged with DNS-discovered IPs."],
    ]
)

add_callout("TIP", "Always provide industry and revenue for the most accurate insurance analytics. Without these, "
            "the financial impact calculator uses conservative defaults that may not reflect the client\u2019s actual risk profile.")

add_callout("WARNING", "Scans take 10\u201318 minutes. Do not close the browser tab during scanning \u2014 the SSE connection "
            "will be lost. If disconnected, you can retrieve results later via the /results/<scan_id> URL or the scan history API.")

add_page_break()

# ============================================================================
# 3. UNDERSTANDING THE RESULTS DASHBOARD
# ============================================================================
add_h1("3. Understanding the Results Dashboard")

add_h2("3.1 Risk Score (0\u20131000)")
add_body(
    "The overall risk score is a weighted composite of all 25 scoring categories, calculated on a 0\u20131000 scale "
    "where higher scores indicate greater risk. The score is divided into four risk bands:"
)
add_table(
    ["Score Range", "Risk Level", "Interpretation"],
    [
        ["0\u2013199", "Low", "Strong external security posture. Minor or no findings. Favourable for underwriting."],
        ["200\u2013399", "Medium", "Some security gaps identified. Addressable issues that should be remediated. Standard underwriting."],
        ["400\u2013599", "High", "Significant security weaknesses. Multiple critical or high-severity findings. Requires remediation before or as a condition of cover."],
        ["600\u20131000", "Critical", "Severe security deficiencies. Exposed databases, unpatched critical CVEs, RDP open to internet, or active compromise indicators. May be uninsurable without immediate remediation."],
    ]
)
add_body(
    "The score is calculated by converting each checker\u2019s result to a 0\u2013100 risk value (higher = riskier), "
    "multiplying by the checker\u2019s weight, summing all weighted risks, and scaling to 0\u20131000. "
    "A WAF bonus of up to 50 points is subtracted if a web application firewall is detected."
)
add_callout("NOTE", "The risk score should be interpreted in context. A score of 350 for a small retail business "
            "has different underwriting implications than the same score for a financial services firm handling "
            "sensitive customer data.")

add_h2("3.2 Executive Summary Cards")
add_body(
    "The dashboard displays a row of summary cards providing at-a-glance metrics. Each card shows a key "
    "indicator with a traffic-light colour (green/amber/red) based on severity:"
)
add_bullet("SSL Grade \u2014 letter grade (A+ to F) for TLS configuration quality. A/A+ is green, B/C is amber, D/F is red.")
add_bullet("Email Score \u2014 score out of 10 for SPF/DKIM/DMARC configuration. 8+ is green, 5\u20137 is amber, below 5 is red.")
add_bullet("Headers % \u2014 percentage of recommended HTTP security headers present. 70%+ is green, 40\u201369% is amber, below 40% is red.")
add_bullet("WAF \u2014 whether a Web Application Firewall was detected. Detected is green, not detected is red.")
add_bullet("Breaches \u2014 number of known data breaches from HIBP. 0 is green, 1\u20133 is amber, 4+ is red.")
add_bullet("Admin Panels \u2014 number of exposed administrative interfaces found. 0 is green, any found is red.")
add_bullet("DB Exposure \u2014 whether database ports (MongoDB, Redis, PostgreSQL, MySQL, etc.) are publicly accessible. None is green, any is red (critical).")
add_bullet("Blacklisting \u2014 whether the domain or IP appears on DNS blacklists. Clean is green, listed is red.")
add_bullet("RDP \u2014 whether Remote Desktop Protocol (port 3389) is exposed to the internet. Not exposed is green, exposed is red (critical).")
add_bullet("Annual Loss \u2014 estimated probable annual loss from the hybrid model (median/P50 value in ZAR).")
add_bullet("Web Ranking \u2014 Tranco list ranking indicating the domain\u2019s relative popularity and visibility.")

add_h2("3.3 Vulnerability Posture")
add_body(
    "This section aggregates all CVE (Common Vulnerabilities and Exposures) data from Shodan and OSV.dev "
    "across all discovered IP addresses. It presents:"
)
add_bullet("CVE severity breakdown \u2014 counts of critical, high, medium, and low severity vulnerabilities.")
add_bullet("CISA KEV matches \u2014 number of CVEs appearing in the CISA Known Exploited Vulnerabilities catalog, confirming active exploitation in the wild.")
add_bullet("EPSS scores \u2014 Exploit Prediction Scoring System probabilities, indicating the likelihood of exploitation within the next 30 days.")
add_bullet("Patch management posture \u2014 age analysis of unpatched CVEs: how many are older than 180 days, 90\u2013180 days, and under 90 days.")
add_bullet("Ransomware associations \u2014 CVEs known to be exploited by ransomware groups.")
add_bullet("MITRE ATT&CK technique mapping \u2014 links CVEs to known attacker techniques and threat groups.")

add_callout("WARNING", "The presence of CISA KEV CVEs is a critical finding. These vulnerabilities have confirmed "
            "active exploitation and should be treated as the highest remediation priority.")

add_h2("3.4 Attacker\u2019s View (Kill Chain)")
add_body(
    "The Attacker\u2019s View maps scan findings to the phases of a typical cyber attack, helping non-technical "
    "users understand how each finding could be exploited. The kill chain phases covered are:"
)
add_bullet("Reconnaissance \u2014 what an attacker can discover about the target (subdomains, technology stack, exposed files, web ranking visibility).")
add_bullet("Initial Access \u2014 entry points an attacker could exploit (RDP, exposed admin panels, weak SSL, phishing via missing DMARC).")
add_bullet("Credential Access \u2014 leaked credentials from breaches and dark web sources (HIBP, Dehashed, Hudson Rock, IntelX).")
add_bullet("Exploitation \u2014 known vulnerabilities that could be exploited (Shodan CVEs, OSV.dev version analysis, exposed databases).")
add_bullet("Impact \u2014 potential consequences (ransomware susceptibility, financial impact, business interruption risk).")

add_h2("3.5 Insurance Analytics")
add_body(
    "The insurance analytics section provides three proprietary indices and a remediation roadmap, "
    "designed specifically for underwriting and broker use:"
)
add_bullet("Ransomware Susceptibility Index (RSI) \u2014 0.0 to 1.0 scale measuring ransomware exposure.")
add_bullet("Data Breach Index (DBI) \u2014 0 to 100 score measuring historical breach exposure quality.")
add_bullet("Financial Impact (Hybrid Model) \u2014 Monte Carlo simulation producing confidence intervals for probable annual loss in ZAR.")
add_bullet("Remediation Roadmap \u2014 prioritised steps with cost estimates and projected annual savings.")
add_body("These are covered in detail in Section 5.")

add_page_break()

# ============================================================================
# 4. DETAILED CHECKER EXPLANATIONS
# ============================================================================
add_h1("4. Detailed Checker Explanations")
add_body(
    "This section describes each checker category in detail: what it checks, why it matters for insurance "
    "underwriting, how scores are calculated, what the data fields mean, and common findings."
)

# 4.1 Discovery
add_h2("4.1 Discovery")

add_h3("IP Discovery")
add_body(
    "The scanner begins by resolving the target domain\u2019s DNS A records to discover all IP addresses "
    "hosting the domain. If client-supplied IPs are provided, they are merged into the discovered set. "
    "After subdomain enumeration, any new IPs resolved from subdomains are also added to the scan scope."
)
add_body("This phase is not scored directly but determines the scope for all IP-level checkers (ports, CVEs, blacklists, protocols).")
add_callout("TIP", "If the client has additional IP addresses (e.g. separate mail servers, office IPs) that are not "
            "linked via DNS, provide them in the Client IPs field to ensure comprehensive coverage.")

add_h3("Web Ranking (Tranco List)")
add_body(
    "Checks the domain\u2019s position in the Tranco top-sites list, a research-grade ranking of the most "
    "popular websites on the internet. This serves as a proxy for the organisation\u2019s online visibility "
    "and attack surface attractiveness."
)
add_body("Scoring: a domain in the top 100,000 scores lower (more visible = more targeted). An unranked domain scores around 30/100 "
         "(slightly risky due to unknown profile). Weight: 2%.")
add_body("Why it matters: highly ranked domains are more likely to be targeted by automated scanners, credential stuffing campaigns, "
         "and brand impersonation attacks.")

# 4.2 Core Security
add_h2("4.2 Core Security")

add_h3("SSL/TLS Certificate")
add_body(
    "Performs a comprehensive assessment of the domain\u2019s SSL/TLS configuration using sslyze (when available) "
    "or Python\u2019s standard library as a fallback. Checks include:"
)
add_bullet("Certificate validity, expiry date, and remaining days until expiry.")
add_bullet("Certificate chain completeness and trust.")
add_bullet("Key size (minimum 2048-bit RSA recommended).")
add_bullet("TLS version support: flags SSL 2.0/3.0 and TLS 1.0/1.1 as deprecated/insecure.")
add_bullet("Cipher suite analysis: identifies weak ciphers (RC4, DES, 3DES, MD5, NULL, EXPORT, ANON).")
add_bullet("HSTS (HTTP Strict Transport Security) header presence.")
add_bullet("OCSP stapling support.")
add_bullet("CAA (Certificate Authority Authorization) DNS records \u2014 controls which CAs can issue certificates.")
add_body("Grading: A+ (95+), A (85\u201394), B (70\u201384), C (55\u201369), D (40\u201354), F (below 40). Deductions are cumulative: "
         "expired certificate = -40, TLS 1.0 = -20, weak cipher = -20, no HSTS = -10, no CAA = -5, etc.")
add_body("Weight: 9% of overall score.")

add_callout("WARNING", "An expired SSL certificate is an immediate red flag for underwriting. It suggests "
            "either poor IT hygiene or abandoned infrastructure, both of which correlate with higher incident rates.")

add_h3("HTTP Security Headers")
add_body(
    "Checks for the presence and quality of HTTP security headers that protect against common web attacks. "
    "The scanner evaluates:"
)
add_bullet("Strict-Transport-Security (HSTS) \u2014 forces HTTPS connections.")
add_bullet("Content-Security-Policy (CSP) \u2014 mitigates XSS attacks. The scanner also evaluates CSP quality: "
           "a policy using 'unsafe-inline' or 'unsafe-eval' is flagged as weak, while a restrictive nonce-based or hash-based policy scores higher.")
add_bullet("X-Frame-Options \u2014 prevents clickjacking.")
add_bullet("X-Content-Type-Options \u2014 prevents MIME type sniffing.")
add_bullet("Permissions-Policy \u2014 controls browser feature access (camera, microphone, geolocation).")
add_bullet("Referrer-Policy \u2014 controls information sent in HTTP referrer headers.")
add_bullet("X-XSS-Protection \u2014 legacy XSS filter (deprecated but still scored).")
add_body("Score is calculated as a percentage of implemented headers out of total recommended headers. Weight: 5%.")

add_h3("WAF / DDoS Protection")
add_body(
    "Detects whether a Web Application Firewall is protecting the domain. Detection methods include: "
    "response header analysis (e.g. Cloudflare\u2019s cf-ray, AWS WAF headers), CNAME/NS inspection, "
    "and fingerprinting of WAF-specific behaviours."
)
add_body(
    "A detected WAF provides a 50-point bonus reduction to the overall risk score (capped). "
    "This reflects the significant protection WAFs provide against automated attacks, SQL injection, XSS, "
    "and volumetric DDoS."
)
add_body("Weight: WAF bonus applied separately (not a weighted category). Detection is binary: detected or not detected.")

add_h3("Website Security")
add_body(
    "Checks foundational website security configurations:"
)
add_bullet("HTTPS enforcement \u2014 whether HTTP requests are redirected to HTTPS (301 redirect).")
add_bullet("Mixed content \u2014 whether the HTTPS page loads insecure HTTP resources.")
add_bullet("Cookie security \u2014 checks for Secure and HttpOnly flags on cookies.")
add_body("Weight: 4% of overall score.")

# 4.3 Information Security
add_h2("4.3 Information Security")

add_h3("Information Disclosure")
add_body(
    "Probes for sensitive files and directories that should not be publicly accessible. The scanner checks "
    "for common misconfigurations that expose:"
)
add_bullet("Environment files (.env) \u2014 often contain database credentials, API keys, and secrets.")
add_bullet("Git repositories (.git/config) \u2014 may expose source code and commit history.")
add_bullet("Backup files (.bak, .sql, .zip) \u2014 may contain database dumps or configuration backups.")
add_bullet("Configuration files (wp-config.php, web.config, etc.) \u2014 application configuration with credentials.")
add_bullet("Debug/status pages (phpinfo.php, server-status, etc.) \u2014 expose server internals.")
add_body("Each exposed path is classified as critical, high, or medium risk. Weight: 5%.")

add_callout("WARNING", "A publicly accessible .env file is one of the most dangerous findings. It typically "
            "contains database passwords, API keys, and encryption secrets that give an attacker immediate access "
            "to backend systems.")

# 4.4 Email Security
add_h2("4.4 Email Security")

add_h3("Email Authentication (SPF/DKIM/DMARC)")
add_body(
    "Evaluates the three pillars of email authentication that prevent domain spoofing and phishing:"
)
add_bullet("SPF (Sender Policy Framework) \u2014 defines which mail servers are authorised to send email on behalf of the domain. "
           "The scanner checks for presence, validity, the dangerous '+all' mechanism, SPF DNS lookup count (RFC 7208 limits to 10), "
           "and redirect usage.")
add_bullet("DKIM (DomainKeys Identified Mail) \u2014 cryptographically signs outbound email. The scanner probes 40+ common "
           "DKIM selectors (Google, Microsoft, SendGrid, Mailchimp, Amazon SES, Everlytic, etc.) to detect configured signing.")
add_bullet("DMARC (Domain-based Message Authentication) \u2014 ties SPF and DKIM together with an enforcement policy. "
           "The scanner evaluates policy level (none/quarantine/reject), pct= (partial enforcement), subdomain policy (sp=), "
           "and aggregate reporting (rua=).")
add_body("Score: 0\u201310, with deductions for missing records or weak policies. Weight: 6%.")
add_body("Scoring breakdown:")
add_bullet("No SPF: -3 points. SPF with '+all': -3 points.")
add_bullet("No DMARC: -4 points. DMARC policy 'none': -2 points. DMARC 'quarantine': -1 point.")
add_bullet("No DKIM selectors found: -2 points.")
add_bullet("Partial DMARC enforcement (pct < 100): -1 point.")

add_callout("TIP", "A DMARC policy of 'reject' with 100% enforcement and aggregate reporting enabled is the gold standard. "
            "This prevents attackers from sending convincing phishing emails that appear to come from the organisation\u2019s domain.")

add_h3("Advanced Email Hardening (MTA-STS, TLS-RPT, DANE, BIMI)")
add_body("Beyond the core SPF/DKIM/DMARC trio, the scanner checks for advanced email security standards:")
add_bullet("MTA-STS (Mail Transfer Agent Strict Transport Security) \u2014 forces TLS encryption for inbound email delivery, "
           "preventing downgrade attacks where an attacker strips encryption from email in transit.")
add_bullet("TLS-RPT (TLS Reporting) \u2014 a DNS TXT record (_smtp._tls) that enables receiving reports about email TLS delivery "
           "failures. Helps organisations monitor whether their email is being delivered securely.")
add_bullet("DANE (DNS-based Authentication of Named Entities) \u2014 uses DNSSEC to bind TLS certificates to DNS, "
           "preventing certificate impersonation. Requires DNSSEC to be enabled.")
add_bullet("BIMI (Brand Indicators for Message Identification) \u2014 allows organisations to display their logo next to "
           "authenticated emails, enhancing brand trust and visual phishing detection.")
add_body("Weight: 2% of overall score.")

# 4.5 Network & Infrastructure
add_h2("4.5 Network & Infrastructure")

add_h3("DNS & Open Ports (including AXFR)")
add_body(
    "Performs comprehensive DNS and port analysis for each discovered IP address:"
)
add_bullet("Open port scanning \u2014 identifies all open TCP ports on each IP address via Shodan InternetDB and full API data.")
add_bullet("Service banner detection \u2014 identifies the software and version running on each open port.")
add_bullet("DNSSEC validation \u2014 checks whether the domain has DNSSEC enabled.")
add_bullet("Reverse DNS \u2014 resolves IP addresses back to hostnames.")
add_bullet("AXFR (Zone Transfer) testing \u2014 attempts a DNS zone transfer against each nameserver. A successful AXFR is a "
           "critical finding: it means the entire DNS zone (all subdomains, mail servers, internal hostnames) is publicly downloadable.")
add_body("Weight: part of the per-IP scoring that feeds into external_ips (3%).")

add_callout("WARNING", "A permitted AXFR zone transfer is a severe misconfiguration. It exposes the organisation\u2019s "
            "entire DNS infrastructure, making reconnaissance trivial for attackers. This should be remediated immediately "
            "by configuring 'allow-transfer { none; };' on all DNS servers.")

add_h3("Database/Service Exposure (High-Risk Protocols)")
add_body(
    "Checks for publicly exposed high-risk services that should never be accessible from the internet:"
)
add_table(
    ["Port", "Service", "Risk Level"],
    [
        ["21", "FTP", "High \u2014 cleartext file transfer"],
        ["23", "Telnet", "Critical \u2014 cleartext remote access"],
        ["25", "SMTP", "Medium \u2014 potential relay"],
        ["110/143", "POP3/IMAP", "High \u2014 cleartext email"],
        ["445", "SMB", "Critical \u2014 ransomware vector"],
        ["1433", "MSSQL", "Critical \u2014 database exposure"],
        ["3306", "MySQL", "Critical \u2014 database exposure"],
        ["3389", "RDP", "Critical \u2014 #1 ransomware vector"],
        ["5432", "PostgreSQL", "Critical \u2014 database exposure"],
        ["5900", "VNC", "Critical \u2014 remote desktop"],
        ["6379", "Redis", "Critical \u2014 often unauthenticated"],
        ["9200", "Elasticsearch", "Critical \u2014 data exposure"],
        ["11211", "Memcached", "High \u2014 DDoS amplification"],
        ["27017", "MongoDB", "Critical \u2014 often unauthenticated"],
    ]
)
add_body("Weight: 8% of overall score. Each exposed critical service adds significant risk to the RSI calculation.")

add_h3("Cloud & CDN Detection")
add_body(
    "Identifies whether the domain uses cloud hosting and/or a Content Delivery Network. CDN usage "
    "indicates DDoS resilience and infrastructure redundancy. The scanner detects major providers "
    "including Cloudflare, AWS CloudFront, Akamai, Fastly, Azure CDN, and Google Cloud CDN."
)

add_h3("VPN / Remote Access")
add_body(
    "Checks for VPN gateway presence (OpenVPN, WireGuard, Cisco, etc.) and critically whether RDP "
    "(port 3389) is exposed to the public internet. RDP exposure is the single strongest ransomware "
    "indicator \u2014 it contributes +0.25 to the RSI (the largest single factor)."
)
add_body("Weight: 4% of overall score.")

add_h3("DNSBL / Blacklists")
add_body(
    "Queries multiple DNS-based blacklists (DNSBLs) to check whether the domain\u2019s IP addresses or "
    "domain name appear on spam, malware, or abuse lists. Blacklisting indicates a history of compromise, "
    "spam distribution, or abuse originating from the organisation\u2019s infrastructure."
)
add_body("Weight: 6% of overall score.")

add_h3("Shodan CVE Analysis")
add_body(
    "Queries the Shodan API for each discovered IP address to retrieve known vulnerabilities (CVEs) "
    "associated with the exposed services and software versions. For each CVE, the scanner retrieves:"
)
add_bullet("CVSS score \u2014 severity rating (0.0\u201310.0).")
add_bullet("EPSS score \u2014 probability of exploitation within 30 days (0.0\u20131.0).")
add_bullet("CISA KEV status \u2014 whether the CVE is in the Known Exploited Vulnerabilities catalog.")
add_bullet("Ransomware association \u2014 whether the CVE is known to be exploited by ransomware groups.")
add_bullet("MITRE ATT&CK mapping \u2014 associated attack techniques and threat groups.")
add_bullet("Exploit maturity \u2014 whether proof-of-concept or weaponised exploits are publicly available.")
add_bullet("CVE age \u2014 days since publication, indicating patch management hygiene.")
add_body("Weight: 7% of overall score.")

# 4.6 Exposure & Reputation
add_h2("4.6 Exposure & Reputation")

add_h3("HIBP Breaches")
add_body(
    "Queries the Have I Been Pwned (HIBP) API for the domain to check whether it has appeared in known "
    "data breaches. Returns breach count, breach dates, and data classes exposed (emails, passwords, "
    "financial data, etc.)."
)
add_body("Scoring: 0 breaches = no risk; each breach adds 15 points of risk (capped at 100). Weight: 7%.")

add_h3("DNSBL / Blacklists")
add_body("(See Section 4.5 above.)")

add_h3("Exposed Admin Panels")
add_body(
    "Probes for common administrative interfaces that should not be publicly accessible. Checks include "
    "WordPress admin (/wp-admin), phpMyAdmin, cPanel, Webmin, database management tools, and various "
    "CMS-specific admin paths. Each found panel is classified by risk level (critical if login page is "
    "accessible, high if redirect detected)."
)
add_body("Weight: 9% of overall score \u2014 one of the highest weights, reflecting the severity of exposed admin access.")

add_h3("Subdomain Exposure (including Takeover Detection)")
add_body(
    "Discovers subdomains through two methods:"
)
add_bullet("Certificate Transparency (CT) logs via crt.sh \u2014 finds all subdomains that have had certificates issued.")
add_bullet("DNS brute-force \u2014 resolves 50+ common prefixes (dev, staging, admin, api, vpn, etc.) to find additional subdomains.")
add_body(
    "Each discovered subdomain is classified as risky if it matches sensitive keywords (dev, staging, admin, "
    "database, backup, etc.). The scanner also performs subdomain takeover detection: it checks each subdomain\u2019s "
    "CNAME record against 30+ known-vulnerable service patterns (GitHub Pages, Heroku, AWS S3, Azure, Shopify, "
    "Netlify, etc.). If the CNAME points to an unclaimed service, it is flagged as a critical takeover vulnerability."
)
add_body("Weight: 2% of overall score.")

add_callout("WARNING", "A subdomain takeover allows an attacker to host content on your domain (e.g. phishing pages "
            "at staging.yourdomain.com). This is particularly dangerous for organisations where domain trust is critical "
            "(financial services, healthcare).")

add_h3("Shodan CVEs")
add_body("(See Section 4.5 above.)")

add_h3("Dehashed (Credential Leak Database)")
add_body(
    "Queries the Dehashed API for leaked credentials associated with the domain. The scanner performs "
    "credential parsing to classify leak severity:"
)
add_bullet("Plaintext passwords \u2014 critical severity, indicating credentials that can be used immediately.")
add_bullet("Hashed passwords \u2014 high severity, potentially crackable depending on hash algorithm.")
add_bullet("Email-only records \u2014 medium severity, useful for phishing but not direct access.")
add_body(
    "Results are enriched with HIBP breach metadata to add breach dates, data classes, and breach source "
    "information. The scanner also performs credential risk classification combining Dehashed, Hudson Rock, "
    "and IntelX data into a unified risk assessment."
)
add_body("Weight: 3% of overall score. Requires Dehashed API key (paid service).")

add_h3("Hudson Rock")
add_body(
    "Queries Hudson Rock\u2019s free API for infostealer-compromised employees and customers. Infostealers "
    "are malware that harvest credentials, cookies, and session tokens from infected machines. Hudson Rock "
    "aggregates data from these infections to identify which organisations have compromised users."
)
add_body("This data feeds into the Credential Risk Assessment.")

add_h3("IntelX (Intelligence X)")
add_body(
    "Queries the Intelligence X API for dark web and leaked data mentions of the domain. IntelX indexes "
    "paste sites, dark web forums, data breach archives, and other underground sources. Results indicate "
    "whether the organisation\u2019s data or credentials have been circulated in criminal ecosystems."
)
add_body("Requires IntelX API key. Weight: feeds into credential risk assessment.")

add_h3("Credential Risk Assessment")
add_body(
    "A meta-checker that aggregates credential exposure data from Dehashed, Hudson Rock, and IntelX "
    "into a unified risk classification. It evaluates:"
)
add_bullet("Total credential exposure volume.")
add_bullet("Credential freshness (recent vs. historical leaks).")
add_bullet("Credential type severity (plaintext vs. hashed vs. email-only).")
add_bullet("Active infostealer infections (from Hudson Rock).")
add_bullet("Dark web presence (from IntelX).")

add_h3("VirusTotal Intelligence")
add_body(
    "Queries the VirusTotal API for the domain to check whether it has been flagged by any of 90+ "
    "security engines as malicious or suspicious. Also retrieves domain categorisation (e.g. phishing, "
    "malware, gambling) and associated domain intelligence."
)
add_body("Weight: 5% of overall score.")

add_h3("Fraudulent Domains (Lookalike/Typosquatting)")
add_body(
    "Generates permutations of the target domain (character substitution, addition, deletion, "
    "homoglyph attacks) and checks whether any resolve to active websites. Lookalike domains are used "
    "for phishing, brand impersonation, and Business Email Compromise (BEC) attacks."
)
add_body("Weight: 4% of overall score. This is a heavyweight check (optional toggle) that adds several minutes to scan time.")

# 4.7 Technology & Governance
add_h2("4.7 Technology & Governance")

add_h3("Technology Stack / EOL Detection")
add_body(
    "Analyses HTTP response headers and page content to identify the technology stack powering the "
    "website. Detects:"
)
add_bullet("Server software \u2014 Apache, nginx, IIS, etc. with version numbers.")
add_bullet("Application framework \u2014 PHP, ASP.NET, Node.js, Python, etc.")
add_bullet("CMS detection \u2014 WordPress, Joomla, Drupal, Wix, Shopify, Squarespace, Magento, PrestaShop.")
add_bullet("JavaScript libraries \u2014 jQuery (with XSS vulnerability check for versions below 3.5.0), AngularJS (EOL).")
add_bullet("End-of-life (EOL) software \u2014 flags software that no longer receives security patches. "
           "The scanner maintains a comprehensive EOL database covering PHP 5.x\u20137.4, Apache 2.2, nginx legacy, "
           "OpenSSL 1.0/1.1.0, Node.js 12\u201316, Python 2, IIS 6\u20137.5, Tomcat 7\u20138.5, and more.")
add_body("Scoring: EOL critical software = -40 points, high = -25, medium = -10. Weight: 5%.")

add_h3("Domain Intelligence")
add_body(
    "Gathers domain registration and infrastructure intelligence: WHOIS data, domain age, registrar, "
    "nameservers, and DNS configuration. Newer domains may indicate higher risk (phishing infrastructure "
    "tends to use recently registered domains)."
)

add_h3("SecurityTrails DNS History")
add_body(
    "Queries the SecurityTrails API for historical DNS data and associated domains. Reveals "
    "infrastructure changes over time, associated domains on shared hosting, and historical IP addresses. "
    "A large number of associated domains on shared infrastructure increases risk due to neighbour compromise."
)
add_body("Weight: 1% of overall score. Requires SecurityTrails API key (free tier available).")

add_h3("Privacy Compliance (POPIA/GDPR)")
add_body(
    "Checks for the presence and completeness of a privacy policy. The scanner evaluates whether "
    "the policy covers required sections for POPIA (Protection of Personal Information Act) and GDPR compliance. "
    "Missing or incomplete privacy policies indicate governance gaps and regulatory risk."
)
add_body("Weight: 2% of overall score.")

add_h3("Security Policy & VDP")
add_body(
    "Checks for the presence of a security.txt file at /.well-known/security.txt (RFC 9116), "
    "which establishes a vulnerability disclosure policy. Also checks for a security page or "
    "responsible disclosure programme. Organisations with a published VDP demonstrate security maturity."
)

add_h3("Payment Security")
add_body(
    "Checks whether the website processes payments and evaluates the security of payment handling:"
)
add_bullet("Detects self-hosted payment forms (collecting card data directly) vs. third-party payment providers.")
add_bullet("Checks whether payment pages use HTTPS.")
add_bullet("Identifies use of PCI-compliant payment processors (Stripe, PayFast, Peach Payments, etc.).")
add_body("A self-hosted payment form is a high-risk finding (score penalty: 80/100 risk). Weight: 2%.")

# 4.8 Compliance Framework Mapping
add_h2("4.8 Compliance Framework Mapping")
add_body(
    "The scanner maps its findings to four major compliance frameworks, providing an indicative "
    "compliance posture based on externally observable controls. The four frameworks are:"
)

add_h3("POPIA (Protection of Personal Information Act)")
add_body("South Africa\u2019s data protection legislation. The scanner maps 10 controls covering encryption in transit (S19a), "
         "security headers (S19b), web application security (S19c), network access control (S19d), email security (S19e), "
         "privacy policy (S20a), data minimisation (S20b), software currency (S21a), breach history (S22a), and credential exposure (S22b).")

add_h3("PCI DSS v4.0")
add_body("Payment Card Industry Data Security Standard. The scanner maps 10 requirements covering default credentials (Req 2a), "
         "system hardening (Req 2b), security policies (Req 2c), TLS configuration (Req 4a), HTTPS enforcement (Req 4b), "
         "patch management (Req 6a), secure coding (Req 6b), payment security (Req 8a), vulnerability scanning (Req 11a), "
         "and threat monitoring (Req 11b).")

add_h3("ISO 27001")
add_body("International information security management standard. The scanner maps 9 controls covering asset inventory (A.8a), "
         "attack surface (A.8b), network security (A.12a), remote access (A.12b), malware/reputation (A.12c), "
         "DDoS resilience (A.12d), encryption standards (A.14a), application security (A.14b), and payment/data handling (A.14c).")

add_h3("NIST CSF 2.0")
add_body("NIST Cybersecurity Framework version 2.0. The scanner maps 12 functions: security policy (GV.1), privacy governance (GV.2), "
         "asset discovery (ID.1), attack surface mapping (ID.2), encryption/TLS (PR.1), security headers/hardening (PR.2), "
         "perimeter defence (PR.3), email authentication (PR.4), vulnerability detection (DE.1), threat intelligence (DE.2), "
         "breach response (RS.1), security disclosure (RS.2), infrastructure resilience (RC.1), and communication recovery (RC.2).")

add_callout("NOTE", "Compliance scoring is based exclusively on externally observable controls. Many compliance "
            "requirements (e.g. access management, logging, incident response procedures, staff training) cannot be "
            "assessed from the outside. The compliance percentages should be interpreted as indicative, not definitive. "
            "Each control is scored pass (70+), partial (40\u201369), or fail (below 40) based on the average score of its "
            "contributing checkers, weighted by control importance.")

add_page_break()

# ============================================================================
# 5. INSURANCE ANALYTICS (DEEP DIVE)
# ============================================================================
add_h1("5. Insurance Analytics (Deep Dive)")

add_h2("5.1 Ransomware Susceptibility Index (RSI)")
add_body(
    "The RSI is a proprietary 0.0\u20131.0 index measuring how susceptible an organisation is to a ransomware "
    "attack, based on externally observable signals. Higher values indicate greater susceptibility."
)
add_table(
    ["RSI Range", "Risk Label", "Interpretation"],
    [
        ["0.00\u20130.24", "Low", "Minimal ransomware exposure. Clean infrastructure, no critical findings."],
        ["0.25\u20130.49", "Medium", "Moderate exposure. Some hygiene issues (missing DMARC, no WAF, credential leaks)."],
        ["0.50\u20130.74", "High", "Significant exposure. Multiple high-risk findings (unpatched CVEs, credential dumps, weak perimeter)."],
        ["0.75\u20131.00", "Critical", "Severe exposure. Critical findings such as exposed RDP, CISA KEV CVEs, and active compromise indicators."],
    ]
)
add_body("The RSI is calculated as follows:")
add_bullet("Base: 0.05 (inherent internet exposure risk for any organisation).")
add_bullet("Priority 1 factors (strongest signals): RDP exposed (+0.25), exposed database ports (+0.10 each, cap 0.20), CISA KEV CVEs (+0.08 each, cap 0.20).")
add_bullet("Priority 2 factors (high impact): high-EPSS CVEs (+0.04 each, cap 0.12), other critical/high CVEs (+0.02 each, cap 0.08), blacklisted IPs (+0.04), critical file exposure (+0.02 each, cap 0.06).")
add_bullet("Priority 3 factors (contributing signals): leaked credentials (scaled by volume: 0.02\u20130.06), breach history (+0.03 if >3 breaches), no DMARC (+0.03), no WAF (+0.03), weak SSL (+0.03).")
add_body("After summing all factors, diminishing returns are applied above 0.50 raw score to prevent score inflation "
         "from stacking many moderate findings. The formula is: if raw <= 0.50, return raw; else 0.50 + 0.50 * (1 - e^(-2*(raw-0.50))).")
add_body("Industry and size multipliers are then applied:")
add_bullet("Industry multipliers (modest, 1.0\u20131.15): healthcare 1.15, legal 1.12, finance/government 1.10\u20131.12, manufacturing/retail/education 1.05, technology/other 1.0.")
add_bullet("Size multipliers: large enterprises (>R500M revenue) get 0.95x, mid-market (>R100M) get 0.98x, SMEs get 1.0x \u2014 reflecting that larger organisations typically have internal defences not visible externally.")

add_callout("TIP", "The RSI is the most influential input to the financial impact calculator. An RSI reduction of 0.10 "
            "through remediation typically translates to a 20\u201340% reduction in projected ransomware losses.")

add_h2("5.2 Data Breach Index (DBI)")
add_body(
    "The DBI scores historical breach exposure on a 0\u2013100 scale (higher = better/less exposed). "
    "It comprises five components:"
)
add_table(
    ["Component", "Max Points", "Scoring"],
    [
        ["Breach Count", "30", "0 breaches = 30 pts, 1\u20133 breaches = 15 pts, 4+ breaches = 0 pts"],
        ["Recency", "20", "No breaches or >3 years ago = 20 pts, 1\u20133 years ago = 10 pts, <1 year ago = 0 pts"],
        ["Data Severity", "15", "No data exposed = 15 pts, emails only = 10 pts, passwords/financials = 0 pts"],
        ["Credential Leaks (Dehashed)", "20", "0 leaks = 20 pts, 1\u2013100 leaks = 10 pts, 100+ leaks = 0 pts, unknown = 10 pts"],
        ["Breach Trend", "15", "No recent breaches (2yr) = 15 pts (Improving), 1\u20132 recent = 7 pts (Stable), 3+ recent = 0 pts (Worsening)"],
    ]
)
add_body("DBI labels: Excellent (80+), Good (60\u201379), Fair (40\u201359), Poor (20\u201339), Critical (below 20).")
add_body("Weight in overall score: 3%.")

add_h2("5.3 Financial Impact Estimation (Hybrid Model)")
add_body(
    "The Financial Impact module uses a hybrid approach derived from FAIR (Factor Analysis of Information Risk) "
    "methodology, anchored to IBM SA 2025 breach cost data (R49.22 million ransom-inclusive average), Sophos SA "
    "2025 ransomware survey data, and actual South African insurance claims data. The model produces four cost "
    "categories, seven incident types, and Monte Carlo confidence intervals for probable annual loss."
)

add_h3("Hybrid engine architecture")
add_body(
    "The total breach magnitude is anchored to the IBM SA average cost (R49.22 million, ransom-inclusive) and "
    "scaled by two factors: revenue scaling with graduated elasticity (smaller organisations experience "
    "proportionally higher costs relative to revenue) and an industry multiplier with graduated severity "
    "(high-risk industries such as financial services, healthcare, and legal receive an uplift reflecting "
    "their higher regulatory exposure and historical claims frequency)."
)
add_body("The anchored magnitude is decomposed into five cost components:")
add_bullet("C1: Post-breach liability (residual) \u2014 notification costs, credit monitoring, legal fees, reputational damage, and customer churn.")
add_bullet("C2: Regulatory fines per jurisdiction \u2014 POPIA fines (up to R10 million), GDPR exposure (4% of global turnover, uncapped) if applicable, PCI DSS fines if applicable.")
add_bullet("C3: Business interruption \u2014 revenue loss during recovery, using SA-calibrated PERT(3, 25, 120) days recovery time per Sophos SA 2025.")
add_bullet("C4: Ransom/extortion \u2014 10.40% of total breach magnitude, proportional to the IBM SA data decomposition. Activated for ransomware-family incidents only.")
add_bullet("C5: Incident response \u2014 forensics, containment, eradication, and recovery. Tiered by organisation size.")

add_h3("Four reporting categories")
add_body("The five cost components are grouped into four reporting categories aligned with standard insurance policy sections:")
add_bullet("Data breach exposure: C1 (post-breach liability) + C2 (regulatory fines).")
add_bullet("Detection and escalation: C5 (incident response).")
add_bullet("Ransom demand: C4 (ransom/extortion).")
add_bullet("Business interruption: C3 (revenue loss during recovery).")

add_h3("Probability model")
add_body(
    "The probability of a breach event is calculated as p_breach = Vulnerability \u00d7 TEF \u00d7 0.30, where "
    "Vulnerability is derived from scan findings and TEF (Threat Event Frequency) is the annual frequency of "
    "attempted attacks for the organisation\u2019s industry and size profile. The 0.30 calibration factor aligns "
    "modelled probabilities with observed SA claims frequencies."
)
add_body(
    "The RSI score drives ransomware-family incidents specifically. Ransomware initial access vector weights "
    "are calibrated to Sophos SA 2025 survey data: compromised credentials 34%, exploited vulnerabilities 28%, "
    "malicious email (phishing) 22%, other vectors 16%."
)

add_h3("Monte Carlo simulation")
add_body(
    "The model runs 10,000 Monte Carlo iterations using PERT distributions (lambda=4). Key PERT parameters "
    "include SA recovery time PERT(3, 25, 120) days, cost-per-record ranges by industry, and ransom demand "
    "distributions calibrated to Sophos SA 2025 median payment data. The output includes:"
)
add_bullet("P5 (5th percentile) \u2014 optimistic scenario, 95% chance the actual loss exceeds this.")
add_bullet("P25 (25th percentile) \u2014 lower quartile.")
add_bullet("P50 (50th percentile / median) \u2014 most likely loss, used as the primary estimate.")
add_bullet("P75 (75th percentile) \u2014 upper quartile.")
add_bullet("P95 (95th percentile) \u2014 pessimistic scenario, only 5% chance the actual loss exceeds this. Used for recommended coverage limit.")
add_body("Insurance recommendations derived from the simulation:")
add_bullet("Suggested deductible: RSI-scaled percentage (0.5%\u201320%) of the recommended coverage limit.")
add_bullet("Expected annual loss: P50 (median).")
add_bullet("Recommended coverage limit: P95 multiplied by 1.2 (20% safety margin).")
add_body("All financial calculations use South African Rand (ZAR) with IBM SA and Sophos SA calibration data.")

add_callout("NOTE", "Financial projections are indicative estimates based on external scanning data and industry averages. "
            "Actual losses depend on many internal factors (backup quality, incident response capability, cyber insurance "
            "coverage, employee awareness) that are not visible in an external scan. Use these figures as a starting point "
            "for underwriting conversations, not as definitive predictions.")

add_h2("5.4 Remediation Roadmap")
add_body(
    "The Remediation Roadmap maps scan findings to prioritised, actionable remediation steps. For each "
    "applicable finding, it provides:"
)
add_bullet("Action \u2014 specific, plain-language instruction for what to fix and why.")
add_bullet("Priority \u2014 1 (critical, immediate action), 2 (high, address within 30 days), or 3 (medium, address within 90 days).")
add_bullet("Estimated cost \u2014 indicative cost range in ZAR for implementing the remediation.")
add_bullet("RSI reduction \u2014 how much the remediation would reduce the Ransomware Susceptibility Index.")
add_bullet("Annual savings estimate \u2014 projected reduction in probable annual loss if the remediation is implemented.")

add_body("Key remediation items and their RSI impact:")
add_table(
    ["Remediation", "Priority", "RSI Reduction", "Estimated Cost"],
    [
        ["Block RDP from public internet", "1 (Critical)", "0.35", "R9,000\u2013R36,000"],
        ["Firewall exposed database ports", "1 (Critical)", "0.15", "R9,000\u2013R36,000"],
        ["Patch CISA KEV CVEs", "1 (Critical)", "0.10", "R18,000\u2013R90,000"],
        ["Patch critical OSV.dev CVEs", "1 (Critical)", "0.10", "R18,000\u2013R90,000"],
        ["Patch high-EPSS CVEs", "2 (High)", "0.05", "R18,000\u2013R90,000"],
        ["Upgrade SSL/TLS", "2 (High)", "0.05", "R0\u2013R3,600"],
        ["Implement DMARC", "2 (High)", "0.05", "R3,600\u2013R9,000"],
        ["Deploy WAF", "2 (High)", "0.05", "R0\u2013R9,000/mo"],
        ["Force credential resets (leaks)", "2 (High)", "0.05", "R9,000\u2013R36,000"],
        ["Resolve blacklisting", "3 (Medium)", "0.05", "R9,000\u2013R36,000"],
        ["Update EOL software", "3 (Medium)", "0.03", "R9,000\u2013R36,000"],
        ["Implement security headers", "2 (High)", "0.03", "R0\u2013R3,600"],
        ["Remove exposed sensitive files", "2 (High)", "0.03", "R0\u2013R9,000"],
    ]
)

add_body(
    "The roadmap also calculates a simulated post-remediation state: the projected RSI and financial impact "
    "if all recommended items were implemented. This \u201cbefore and after\u201d comparison is powerful for broker "
    "presentations, showing clients the tangible financial benefit of security improvements."
)

add_callout("WARNING", "Cost estimates are indicative ranges based on South African IT service rates. Actual costs "
            "will vary based on the organisation\u2019s existing infrastructure, IT vendor, and complexity of implementation.")

add_page_break()

# ============================================================================
# 6. PDF REPORTS
# ============================================================================
add_h1("6. PDF Reports")

add_h2("6.1 Full Technical Report")
add_body(
    "The full technical PDF report is the comprehensive output of a scan, intended for technical staff "
    "and detailed underwriting review. It includes:"
)
add_bullet("Cover page with domain, scan date, risk score, and risk level.")
add_bullet("Executive summary with all summary cards and key metrics.")
add_bullet("Vulnerability posture with full CVE table, severity breakdown, and patch management analysis.")
add_bullet("Detailed findings for every checker category with specific issues and recommendations.")
add_bullet("Insurance analytics: RSI, DBI, Financial Impact with Monte Carlo confidence intervals.")
add_bullet("Compliance framework mapping for POPIA, PCI DSS, ISO 27001, and NIST CSF 2.0.")
add_bullet("Remediation roadmap with cost estimates and projected savings.")
add_bullet("Technical appendix with all discovered IPs, ports, subdomains, and raw data.")

add_h2("6.2 Broker Summary")
add_body(
    "The broker summary is a condensed version of the report designed for non-technical audiences. "
    "It focuses on:"
)
add_bullet("Overall risk score and level with plain-language interpretation.")
add_bullet("Key findings \u2014 top 5\u201310 most impactful issues in non-technical language.")
add_bullet("Financial impact summary \u2014 estimated annual loss range and recommended coverage.")
add_bullet("Remediation highlights \u2014 top 3\u20135 actions the client should take, with cost estimates.")
add_body("Technical details (CVE lists, port numbers, protocol names) are omitted or simplified.")

add_h2("6.3 Reading the Report")
add_body("Reports use the following visual conventions:")
add_bullet("Traffic light colours \u2014 green (good/low risk), amber/orange (moderate/needs attention), red (poor/high risk).")
add_bullet("Card layout \u2014 each finding category is presented in a card with a header, score, and details.")
add_bullet("Narrative text \u2014 each section includes a plain-language narrative explaining the findings and their implications for insurance.")
add_bullet("Recommendations \u2014 appear as action items with priority levels (Critical/High/Medium).")
add_bullet("Severity badges \u2014 Critical (red), High (orange), Medium (yellow), Low (blue), Info (grey).")

add_page_break()

# ============================================================================
# 7. SCORING METHODOLOGY
# ============================================================================
add_h1("7. Scoring Methodology")

add_h2("7.1 Overall Risk Score")
add_body(
    "The overall risk score is a weighted sum of individual category risk values, scaled to 0\u20131000. "
    "Each category produces a risk value from 0\u2013100 (where 100 = maximum risk). The weighted sum is "
    "multiplied by 10 to produce the final 0\u20131000 score."
)
add_body("The complete weight table:")
add_table(
    ["Category", "Weight", "Description"],
    [
        ["SSL/TLS", "9%", "Certificate validity, TLS versions, cipher strength, HSTS, CAA"],
        ["Exposed Admin Panels", "9%", "Publicly accessible admin interfaces"],
        ["High-Risk Protocols", "8%", "Exposed database and service ports"],
        ["HIBP Breaches", "7%", "Known data breach history"],
        ["Shodan CVEs", "7%", "Known vulnerabilities on external IPs"],
        ["Email Authentication", "6%", "SPF, DKIM, DMARC configuration"],
        ["DNSBL Blacklisting", "6%", "IP/domain blacklist presence"],
        ["Ransomware Risk (RSI)", "6%", "Ransomware susceptibility index"],
        ["HTTP Security Headers", "5%", "Security header implementation"],
        ["Technology Stack / EOL", "5%", "End-of-life software detection"],
        ["VirusTotal Intelligence", "5%", "Malicious/suspicious domain flags"],
        ["Information Disclosure", "5%", "Exposed sensitive files and paths"],
        ["Website Security", "4%", "HTTPS enforcement, mixed content, cookies"],
        ["VPN / Remote Access", "4%", "VPN presence, RDP exposure"],
        ["Fraudulent Domains", "4%", "Lookalike/typosquatting domains"],
        ["Dehashed Credentials", "3%", "Leaked credentials in breach databases"],
        ["External IP Risk", "3%", "Aggregate CVE risk across all IPs"],
        ["Data Breach Index", "3%", "Historical breach quality score"],
        ["Email Hardening", "2%", "MTA-STS, TLS-RPT, DANE, BIMI"],
        ["Payment Security", "2%", "Payment form security"],
        ["Subdomains", "2%", "Risky subdomain exposure"],
        ["Privacy Compliance", "2%", "Privacy policy completeness"],
        ["Web Ranking", "2%", "Domain visibility/popularity"],
        ["Financial Impact", "2%", "Hybrid model loss estimate"],
        ["SecurityTrails DNS", "1%", "Historical DNS and infrastructure"],
    ]
)

add_callout("NOTE", "Weights sum to approximately 100%. The WAF bonus is applied separately as a score reduction "
            "(up to 50 points) rather than as a weighted category.")

add_h2("7.2 Scoring Failsafe")
add_body(
    "When a checker fails (errors or times out during execution), the scoring engine automatically "
    "redistributes its weight proportionally across all remaining successful checkers. This prevents "
    "a failed checker from artificially inflating or deflating the risk score."
)
add_body("The failsafe distinguishes between two types of non-scoring checkers:")
add_bullet("Failed (error/timeout) \u2014 the checker ran but encountered a technical problem. These trigger a scan completeness warning.")
add_bullet("Skipped (no_api_key/disabled/auth_failed) \u2014 the checker was intentionally not run (e.g. no API key configured, toggle off). These do not trigger a warning.")
add_body(
    "Weight redistribution formula: for each remaining checker, its effective weight is multiplied by "
    "(total_weight / remaining_weight) to ensure the effective weights still sum to 100%."
)

add_h2("7.3 Scan Completeness")
add_body(
    "After each scan, the engine reports scan completeness as a percentage:"
)
add_body("Completeness % = (1 - failed_count / assessable_checkers) x 100")
add_body(
    "Where assessable_checkers = total checkers minus intentionally skipped ones. A completeness of 100% "
    "means all enabled checkers ran successfully. The completeness percentage and list of failed checkers "
    "are displayed in both the dashboard and PDF reports."
)

add_h2("7.4 Score Reliability")
add_body("The score is considered reliable when:")
add_bullet("Scan completeness is 100% (no checker failures).")
add_bullet("Key API-dependent checkers (Shodan, HIBP) completed successfully.")
add_bullet("The target domain was accessible during the scan (not down or blocking requests).")
add_body(
    "If completeness is below 100%, a re-scan is recommended. Common causes of checker failure include "
    "target site rate limiting, API rate limits, DNS resolution failures, and network timeouts."
)

add_callout("TIP", "For underwriting decisions, always check the scan completeness percentage. If it is below 90%, "
            "the risk score may be significantly less reliable and a re-scan should be performed.")

add_page_break()

# ============================================================================
# 8. API INTEGRATIONS & DATA SOURCES
# ============================================================================
add_h1("8. API Integrations & Data Sources")
add_body(
    "The scanner integrates with 10 external APIs and data sources to enrich its findings. "
    "The table below summarises each integration:"
)
add_table(
    ["API / Source", "What It Provides", "Cost", "Toggle / Always-On"],
    [
        ["Shodan", "Open ports, services, CVEs, banners, ASN/org data for each IP. Uses InternetDB (free) and full API (key required for deep data).", "Free tier / Paid", "Always-on (InternetDB); full API needs key"],
        ["VirusTotal", "Domain malicious/suspicious flags from 90+ security engines, domain categorisation, associated domains.", "Free tier (4 req/min)", "Always-on (with key)"],
        ["SecurityTrails", "Historical DNS records, associated domains, infrastructure changes over time.", "Free tier (50 req/mo)", "Always-on (with key)"],
        ["Dehashed", "Leaked credentials (emails, passwords, hashes) from breach databases.", "Paid ($5/week)", "Optional toggle"],
        ["Hudson Rock", "Infostealer-compromised employees and customers. Free cavalier API.", "Free", "Always-on"],
        ["Intelligence X (IntelX)", "Dark web mentions, paste sites, leaked data archives.", "Free tier / Paid", "Optional toggle"],
        ["HIBP (Have I Been Pwned)", "Domain breach history: breach names, dates, data classes, affected record counts.", "Free (domain search) / Paid (email search)", "Always-on"],
        ["CISA KEV", "Known Exploited Vulnerabilities catalog \u2014 CVEs with confirmed active exploitation.", "Free (public catalog)", "Always-on"],
        ["FIRST.org EPSS", "Exploit Prediction Scoring System \u2014 30-day exploitation probability for each CVE.", "Free (public API)", "Always-on"],
        ["OSV.dev", "Open Source Vulnerabilities database \u2014 maps CPE/package versions to known CVEs for version-based enrichment.", "Free (public API)", "Always-on"],
    ]
)

add_callout("NOTE", "API keys are configured via environment variables in the .env file. The scanner functions "
            "without any API keys but produces less comprehensive results. For production underwriting use, "
            "Shodan, HIBP, and VirusTotal keys are strongly recommended at minimum.")

add_page_break()

# ============================================================================
# 9. KNOWN LIMITATIONS & PLANNED IMPROVEMENTS
# ============================================================================
add_h1("9. Known Limitations & Planned Improvements")

add_h2("Current Limitations")
add_bullet("External-only scanning \u2014 cannot assess internal controls, endpoint protection, backup policies, MFA, or employee awareness training.")
add_bullet("OSV.dev API reliability \u2014 the OSV.dev API can be intermittently slow or unavailable, which may result in incomplete CVE enrichment from version analysis.")
add_bullet("Scan speed \u2014 current scan time is 10\u201318 minutes depending on target complexity and API response times. An optimisation roadmap targets reducing this to 5\u20138 minutes through parallel execution improvements and caching.")
add_bullet("Rate limiting on target sites \u2014 aggressive web application firewalls or rate limiters on the target domain may block some probes, resulting in incomplete data for information disclosure, admin panel, and website security checks.")
add_bullet("CDN/WAF masking \u2014 domains behind CDNs may show CDN IP addresses rather than origin server IPs, potentially missing vulnerabilities on the actual hosting infrastructure.")
add_bullet("Fraudulent domain scanning \u2014 the lookalike domain checker generates a finite set of permutations and cannot detect all possible typosquatting variants. Sophisticated homoglyph attacks using Unicode characters may be missed.")
add_bullet("Shodan data freshness \u2014 Shodan\u2019s InternetDB is updated periodically (not real-time). Recent infrastructure changes may not yet be reflected in results.")
add_bullet("Single-scan perspective \u2014 the scanner provides a point-in-time snapshot. Continuous monitoring is not yet implemented.")

add_h2("Planned Improvements")
add_body("The following improvements are on the development roadmap (refer to the Gap Analysis document for the full plan):")
add_bullet("Continuous monitoring \u2014 scheduled re-scans with delta comparison and alerting on security posture changes.")
add_bullet("Port depth improvements \u2014 active port scanning alongside Shodan data for more complete port coverage.")
add_bullet("Technology fingerprinting \u2014 deeper CMS version detection and framework identification using multiple detection methods.")
add_bullet("Broker API integration \u2014 REST API for programmatic scan initiation and result retrieval by broker systems.")
add_bullet("Per-checker accuracy improvements \u2014 individual accuracy enhancements for all 22+ security checkers based on false positive/negative analysis.")
add_bullet("Dark web monitoring alternatives \u2014 evaluation of alternative dark web APIs beyond IntelX for cost-effective credential monitoring.")
add_bullet("Dynamic CVE re-prioritisation \u2014 integration with threat feed data (e.g. Sophos) for real-time CVE priority adjustments based on active exploitation campaigns.")

add_page_break()

# ============================================================================
# 10. GLOSSARY
# ============================================================================
add_h1("10. Glossary")

glossary = [
    ("AXFR", "DNS zone transfer protocol. A successful AXFR query returns the complete contents of a DNS zone (all records). Should be restricted to authorised secondary nameservers only."),
    ("BIMI", "Brand Indicators for Message Identification. An email standard that allows senders to display a brand logo next to authenticated messages in supported email clients."),
    ("CAA", "Certificate Authority Authorization. DNS records that specify which certificate authorities are permitted to issue SSL/TLS certificates for a domain."),
    ("CDN", "Content Delivery Network. A distributed network of servers that cache and deliver web content from locations closer to the end user, improving performance and DDoS resilience."),
    ("CSP", "Content Security Policy. An HTTP header that controls which resources (scripts, styles, images) a browser is allowed to load, mitigating cross-site scripting (XSS) attacks."),
    ("CVE", "Common Vulnerabilities and Exposures. A standardised identifier for publicly known security vulnerabilities (e.g. CVE-2024-12345)."),
    ("CVSS", "Common Vulnerability Scoring System. A numerical score (0.0\u201310.0) indicating the severity of a vulnerability. Critical: 9.0\u201310.0, High: 7.0\u20138.9, Medium: 4.0\u20136.9, Low: 0.1\u20133.9."),
    ("DANE", "DNS-based Authentication of Named Entities. Uses DNSSEC to bind TLS certificates to DNS records, preventing certificate impersonation."),
    ("DBI", "Data Breach Index. A Phishield proprietary score (0\u2013100) measuring historical breach exposure quality. Higher is better (less exposed)."),
    ("DKIM", "DomainKeys Identified Mail. An email authentication method that uses cryptographic signatures to verify that an email was sent by an authorised server and has not been tampered with."),
    ("DMARC", "Domain-based Message Authentication, Reporting and Conformance. An email authentication protocol that uses SPF and DKIM to prevent domain spoofing, with enforcement policies (none/quarantine/reject)."),
    ("DNSBL", "DNS-based Blacklist. Lists of IP addresses or domains known to be associated with spam, malware, or abuse. Email servers query DNSBLs to filter incoming mail."),
    ("DNSSEC", "Domain Name System Security Extensions. Adds cryptographic signatures to DNS records to prevent DNS spoofing and cache poisoning attacks."),
    ("EPSS", "Exploit Prediction Scoring System. A probability score (0.0\u20131.0) from FIRST.org predicting the likelihood that a CVE will be exploited in the wild within the next 30 days."),
    ("Hybrid Financial Impact Model (derived from FAIR)", "A quantitative risk analysis model that expresses cyber risk in financial terms. The Phishield scanner uses a hybrid approach derived from FAIR principles, anchored to IBM SA breach cost data and calibrated with Sophos SA 2025 and actual insurance claims data."),
    ("HSTS", "HTTP Strict Transport Security. An HTTP header that forces browsers to use HTTPS for all future connections to a domain, preventing protocol downgrade attacks."),
    ("KEV", "Known Exploited Vulnerabilities. A CISA-maintained catalog of CVEs with confirmed active exploitation in the wild. Federal agencies must patch KEV entries within mandated timeframes."),
    ("Kill Chain", "A model describing the stages of a cyber attack: Reconnaissance, Initial Access, Credential Access, Exploitation, and Impact."),
    ("Monte Carlo", "A statistical simulation method that runs thousands of random iterations to model uncertainty and produce probability distributions for financial outcomes."),
    ("MTA-STS", "Mail Transfer Agent Strict Transport Security. Forces TLS encryption for inbound email delivery by publishing a policy via HTTPS, preventing email encryption downgrade attacks."),
    ("NIST CSF", "National Institute of Standards and Technology Cybersecurity Framework. A widely adopted framework for managing cybersecurity risk, organised into six functions: Govern, Identify, Protect, Detect, Respond, Recover."),
    ("PCI DSS", "Payment Card Industry Data Security Standard. A set of security standards for organisations that handle credit card data, covering network security, access control, and monitoring."),
    ("PERT", "Program Evaluation and Review Technique. A probability distribution used in Monte Carlo simulations that concentrates more probability around the most likely value than a triangular distribution."),
    ("POPIA", "Protection of Personal Information Act. South Africa\u2019s data protection legislation, similar to GDPR, governing how organisations collect, process, and store personal information."),
    ("RDP", "Remote Desktop Protocol. Microsoft\u2019s protocol for remote desktop access (port 3389). When exposed to the internet, it is the single most common ransomware entry vector."),
    ("RSI", "Ransomware Susceptibility Index. A Phishield proprietary score (0.0\u20131.0) measuring how susceptible an organisation is to ransomware based on externally observable signals."),
    ("SPF", "Sender Policy Framework. A DNS-based email authentication mechanism that specifies which mail servers are authorised to send email on behalf of a domain."),
    ("SSE", "Server-Sent Events. A web technology enabling the server to push real-time updates to the browser over a persistent HTTP connection. Used by the scanner for live progress tracking."),
    ("TLS", "Transport Layer Security. The cryptographic protocol that provides encryption for data in transit. TLS 1.2 and 1.3 are the current recommended versions."),
    ("TLS-RPT", "TLS Reporting. A DNS TXT record (_smtp._tls) that enables organisations to receive reports about email TLS delivery failures, helping monitor the effectiveness of MTA-STS deployment."),
    ("VDP", "Vulnerability Disclosure Policy. A published process for security researchers to report vulnerabilities, typically via a security.txt file at /.well-known/security.txt."),
    ("WAF", "Web Application Firewall. A security layer that filters and monitors HTTP traffic to protect web applications from attacks such as SQL injection, XSS, and DDoS."),
    ("ZTNA", "Zero Trust Network Access. A security model that requires verification for every person and device trying to access resources, regardless of network location. An alternative to traditional VPN."),
]

add_table(
    ["Term", "Definition"],
    [[term, defn] for term, defn in glossary]
)

add_page_break()

# ============================================================================
# 11. VERSION HISTORY
# ============================================================================
add_h1("11. Version History")
add_body("This document is maintained as a living reference and will be updated as new features are added to the scanner.")

add_table(
    ["Version", "Date", "Changes"],
    [
        ["1.0", "April 2026", "Initial release covering all current scanner capabilities: 25+ checkers across 8 categories, "
         "10 API integrations, RSI/DBI/hybrid financial impact analytics, Monte Carlo simulation, compliance framework mapping "
         "(POPIA, PCI DSS, ISO 27001, NIST CSF 2.0), subdomain takeover detection, AXFR testing, CSP quality analysis, "
         "TLS-RPT checking, credential parsing, CAA records, and remediation roadmap with cost estimates."],
    ]
)

add_body("")
add_body("")
p = doc.add_paragraph()
r = p.add_run("--- End of Document ---")
r.italic = True
r.font.color.rgb = RGBColor(128, 128, 128)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# SAVE
# ============================================================================
doc.save(OUTPUT)
print(f"Manual saved to: {OUTPUT}")
