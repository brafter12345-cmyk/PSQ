"""Convert OUTSTANDING.md to OUTSTANDING.docx for easy reading.

The .md file is the live working copy (edited frequently, git-diff-able).
The .docx file is a polished read-only snapshot regenerated on demand
via this script.

Usage (from security_scanner/ directory):
    py -3 tooling/generate_outstanding_docx.py

Output: security_scanner/docs/OUTSTANDING.docx

Re-run whenever OUTSTANDING.md changes if you want the docx kept in sync.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
ROOT = HERE.parent
SRC = ROOT / "docs" / "OUTSTANDING.md"
OUT = ROOT / "docs" / "OUTSTANDING.docx"

# Phishield brand colors (match the gap analysis + PDF report)
NAVY = RGBColor(0x0F, 0x27, 0x44)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
GREY_DARK = RGBColor(0x47, 0x55, 0x69)
GREY_MID = RGBColor(0x64, 0x74, 0x8B)


def _set_cell_bg(cell, hex_color):
    """Set table cell background color via raw OOXML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _apply_inline_formatting(paragraph, text):
    """Parse inline markdown (**bold**, *italic*, `code`, [text](link)) and
    append runs to the paragraph with appropriate formatting."""
    # Regex covers bold (**...**), italic (*...*), inline code (`...`),
    # and links ([text](url)). Order matters: try bold/code/link before
    # italic so `**word**` isn't parsed as italic-bold-italic.
    pattern = re.compile(
        r'(\*\*[^\*]+\*\*)|'
        r'(`[^`]+`)|'
        r'(\[[^\]]+\]\([^\)]+\))|'
        r'(\*[^\*]+\*)'
    )
    pos = 0
    for m in pattern.finditer(text):
        # plain text before the match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('`') and token.endswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        elif token.startswith('[') and ']' in token and '(' in token:
            # [display](url) - render as 'display (url)' since this is
            # an internal doc and full hyperlinking is overkill
            display, _, rest = token[1:].partition(']')
            url = rest.lstrip('(').rstrip(')')
            r = paragraph.add_run(display)
            r.font.color.rgb = BLUE
            r.underline = True
            # Append the URL in plain text after the display name if it
            # adds useful information (not just a duplicate)
            if url not in display:
                paragraph.add_run(f" ({url})").font.size = Pt(9)
        elif token.startswith('*') and token.endswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    # remainder
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_heading(doc, text, level):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    if level == 1:
        run.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def _set_fixed_layout(table):
    """Force w:tblLayout type=fixed so Word honours the column widths instead
    of growing wide columns past the right margin (autofit caused a right-side
    cut-off on tables with long tokens)."""
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)


def _add_table(doc, rows):
    """Add a markdown-style table. `rows` is a list of row-cells; first
    row is treated as header."""
    if not rows or not rows[0]:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Fixed layout + explicit even column widths: keeps the table within the
    # page's content area so a long token cannot push a column off the right
    # edge. python-docx needs the width set on every cell for fixed layout.
    table.autofit = False
    _set_fixed_layout(table)
    sec = doc.sections[0]
    col_w = int((sec.page_width - sec.left_margin - sec.right_margin) / n_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_w
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row[:n_cols]):
            cell = table.rows[r_idx].cells[c_idx]
            # Clear default empty paragraph
            cell.text = ""
            p = cell.paragraphs[0]
            _apply_inline_formatting(p, cell_text.strip())
            # Style: header row navy bg + white text
            if r_idx == 0:
                _set_cell_bg(cell, '0F2744')
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(9.5)
            else:
                if r_idx % 2 == 0:
                    _set_cell_bg(cell, 'F2F7FA')
                for r in p.runs:
                    r.font.size = Pt(9.5)


def _parse_md(md_text):
    """Yield rendering ops from the markdown source. Each op is a tuple:
       ('h1', text) / ('h2', text) / ('h3', text)
       ('p', text)
       ('bullet', text)
       ('table', [[cell, cell, ...], ...])
       ('hr',)
    """
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Headings
        if line.startswith('# '):
            yield ('h1', line[2:].strip())
        elif line.startswith('## '):
            yield ('h2', line[3:].strip())
        elif line.startswith('### '):
            yield ('h3', line[4:].strip())
        elif line.strip() == '---':
            yield ('hr',)
        # Tables: a row starts with '|' and the next line is the
        # divider with hyphens. Consume the block.
        elif line.startswith('|') and i + 1 < len(lines) and re.match(
                r'^\|\s*[:\-]+', lines[i + 1].strip()):
            rows = []
            # header row
            header = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(header)
            i += 2  # skip header + divider
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            yield ('table', rows)
            continue
        # Bullet list
        elif re.match(r'^\s*[-*]\s', line):
            yield ('bullet', re.sub(r'^\s*[-*]\s', '', line))
        elif line.strip().startswith('1.') or re.match(r'^\d+\.\s', line):
            # Numbered list - treat as bullet for simplicity
            yield ('bullet', re.sub(r'^\d+\.\s+', '', line))
        elif line.strip() == '':
            yield ('blank',)
        else:
            yield ('p', line)
        i += 1


def build_docx():
    if not SRC.exists():
        print(f"Source not found: {SRC}", file=sys.stderr)
        sys.exit(1)
    md_text = SRC.read_text(encoding='utf-8')
    doc = Document()
    # Default font + margins for a clean read
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Document title (Phishield-style header)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHISHIELD CYBER RISK SCANNER")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Outstanding Items")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(20)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Generated from OUTSTANDING.md")
    r.italic = True
    r.font.color.rgb = GREY_MID
    r.font.size = Pt(9)

    doc.add_paragraph()

    for op in _parse_md(md_text):
        kind = op[0]
        if kind == 'h1':
            # The .md starts with '# Outstanding Items' which we already
            # rendered as the title page; skip it.
            if op[1].lower().startswith('outstanding'):
                continue
            _add_heading(doc, op[1], 1)
        elif kind == 'h2':
            _add_heading(doc, op[1], 2)
        elif kind == 'h3':
            _add_heading(doc, op[1], 3)
        elif kind == 'p':
            p = doc.add_paragraph()
            _apply_inline_formatting(p, op[1])
        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            _apply_inline_formatting(p, op[1])
        elif kind == 'table':
            _add_table(doc, op[1])
            doc.add_paragraph()  # spacer after each table
        elif kind == 'hr':
            # Horizontal rule - represent as a centered separator
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('— — —')
            r.font.color.rgb = GREY_MID
            r.font.size = Pt(9)
        elif kind == 'blank':
            # Skip - paragraphs already insert their own spacing
            pass

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Generated: {OUT}")
    print(f"Size: {OUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    build_docx()
